# portfolio_manager_fixed.py
import os
import sys
import sqlite3
import datetime
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import matplotlib.pyplot as plt
import matplotlib
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import webbrowser

# Используем агрессивный бэкэнд для matplotlib чтобы избежать проблем с GUI
matplotlib.use('Agg')


class PortfolioDatabase:
    """Класс для работы с базой данных"""

    def __init__(self, db_path='portfolio.db'):
        self.db_path = db_path
        self.init_db()

    def get_connection(self):
        """Создает соединение с базой данных"""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn

    def init_db(self):
        """Инициализирует базу данных с необходимыми таблицами"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            # Таблица записей
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS records (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT NOT NULL,
                    type TEXT NOT NULL,
                    year INTEGER NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    file_path TEXT UNIQUE
                )
            ''')

            # Таблица соавторов
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS coauthors (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    record_id INTEGER NOT NULL,
                    name TEXT NOT NULL,
                    FOREIGN KEY (record_id) REFERENCES records (id) ON DELETE CASCADE,
                    UNIQUE(record_id, name)
                )
            ''')

            # Таблица логов активности
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS activity_log (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    record_id INTEGER,
                    action TEXT NOT NULL,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    details TEXT
                )
            ''')

            conn.commit()

    def create_record(self, title, record_type, year):
        """Создает новую запись в базе данных"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            # Создаем имя файла
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '_')
            file_name = f"{safe_title}_{year}.md"

            # Создаем директорию если нужно
            base_dir = os.path.dirname(os.path.abspath(__file__))
            records_dir = os.path.join(base_dir, 'records')
            os.makedirs(records_dir, exist_ok=True)

            # Абсолютный путь к файлу
            file_path = os.path.join(records_dir, file_name)

            # Относительный путь для БД
            relative_path = os.path.join('records', file_name)

            # Создаем пустой markdown файл
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"# {title}\n\nОписание записи...")

            # Вставляем запись в БД
            cursor.execute('''
                INSERT INTO records (title, type, year, file_path)
                VALUES (?, ?, ?, ?)
            ''', (title, record_type, year, relative_path))

            record_id = cursor.lastrowid

            # Логируем действие
            cursor.execute('''
                INSERT INTO activity_log (record_id, action, details)
                VALUES (?, ?, ?)
            ''', (record_id, 'create', f'Создана запись: {title}'))

            conn.commit()

            return record_id

    def get_all_records(self):
        """Возвращает все записи из базы данных"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT r.*, 
                       GROUP_CONCAT(c.name, ', ') as coauthors
                FROM records r
                LEFT JOIN coauthors c ON r.id = c.record_id
                GROUP BY r.id
                ORDER BY r.created_at DESC
            ''')
            records = cursor.fetchall()

            # Конвертируем Row объекты в словари
            result = []
            for record in records:
                record_dict = dict(record)
                # Получаем абсолютный путь к файлу
                if record_dict['file_path']:
                    record_dict['abs_file_path'] = self.get_absolute_file_path(record_dict['file_path'])
                else:
                    record_dict['abs_file_path'] = None
                result.append(record_dict)

            return result

    def get_record(self, record_id):
        """Возвращает запись по ID"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT r.*, 
                       GROUP_CONCAT(c.name, ', ') as coauthors
                FROM records r
                LEFT JOIN coauthors c ON r.id = c.record_id
                WHERE r.id = ?
                GROUP BY r.id
            ''', (record_id,))

            record = cursor.fetchone()
            if record:
                # Конвертируем Row в словарь
                record_dict = dict(record)

                # Получаем абсолютный путь
                if record_dict.get('file_path'):
                    record_dict['abs_file_path'] = self.get_absolute_file_path(record_dict['file_path'])
                else:
                    record_dict['abs_file_path'] = None

                return record_dict
            return None

    def update_record(self, record_id, title, record_type, year, content):
        """Обновляет запись"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            # Получаем старую запись
            cursor.execute('SELECT file_path FROM records WHERE id = ?', (record_id,))
            old_record = cursor.fetchone()
            if not old_record:
                return False

            old_relative_path = dict(old_record)['file_path']

            # Получаем абсолютные пути
            base_dir = os.path.dirname(os.path.abspath(__file__))
            old_abs_path = os.path.join(base_dir, old_relative_path) if old_relative_path else None

            # Создаем новое имя файла
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '_')
            new_file_name = f"{safe_title}_{year}.md"

            # Создаем директорию если нужно
            records_dir = os.path.join(base_dir, 'records')
            os.makedirs(records_dir, exist_ok=True)

            # Новые пути
            new_abs_path = os.path.join(records_dir, new_file_name)
            new_relative_path = os.path.join('records', new_file_name)

            # Переименовываем файл если нужно
            if old_abs_path and os.path.exists(old_abs_path) and old_abs_path != new_abs_path:
                os.rename(old_abs_path, new_abs_path)

            # Обновляем запись в БД
            cursor.execute('''
                UPDATE records 
                SET title = ?, type = ?, year = ?, file_path = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            ''', (title, record_type, year, new_relative_path, record_id))

            # Обновляем содержимое файла
            with open(new_abs_path, 'w', encoding='utf-8') as f:
                f.write(content)

            # Логируем действие
            cursor.execute('''
                INSERT INTO activity_log (record_id, action, details)
                VALUES (?, ?, ?)
            ''', (record_id, 'update', f'Обновлена запись: {title}'))

            conn.commit()
            return True

    def get_absolute_file_path(self, relative_path):
        """Конвертирует относительный путь в абсолютный"""
        if not relative_path:
            return None
        base_dir = os.path.dirname(os.path.abspath(__file__))
        return os.path.join(base_dir, relative_path)

    def delete_record(self, record_id):
        """Удаляет запись"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            # Получаем информацию о записи
            record = self.get_record(record_id)
            if not record:
                return False

            # Удаляем файл
            file_path = record.get('abs_file_path')
            if file_path and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Ошибка при удалении файла {file_path}: {e}")

            # Логируем действие
            cursor.execute('''
                INSERT INTO activity_log (record_id, action, details)
                VALUES (?, ?, ?)
            ''', (record_id, 'delete', f'Удалена запись: {record["title"]}'))

            # Удаляем запись из БД (каскадно удалятся соавторы)
            cursor.execute('DELETE FROM records WHERE id = ?', (record_id,))

            conn.commit()
            return True

    def add_coauthor(self, record_id, name):
        """Добавляет соавтора к записи"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            try:
                cursor.execute('''
                    INSERT INTO coauthors (record_id, name)
                    VALUES (?, ?)
                ''', (record_id, name.strip()))
                conn.commit()
                return True
            except sqlite3.IntegrityError:
                # Соавтор уже существует
                return False

    def remove_coauthor(self, record_id, name):
        """Удаляет соавтора из записи"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                DELETE FROM coauthors 
                WHERE record_id = ? AND name = ?
            ''', (record_id, name.strip()))
            conn.commit()

    def get_coauthors(self, record_id):
        """Возвращает соавторов записи"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT name FROM coauthors 
                WHERE record_id = ? 
                ORDER BY name
            ''', (record_id,))
            return [row['name'] for row in cursor.fetchall()]

    def get_statistics(self):
        """Собирает статистику для отчетов"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            # Общее количество записей
            cursor.execute('SELECT COUNT(*) as count FROM records')
            total_records = cursor.fetchone()['count']

            # Распределение по типам
            cursor.execute('''
                SELECT type, COUNT(*) as count 
                FROM records 
                GROUP BY type 
                ORDER BY count DESC
            ''')
            type_distribution = {}
            for row in cursor.fetchall():
                type_distribution[row['type']] = row['count']

            # Распределение по годам
            cursor.execute('''
                SELECT year, COUNT(*) as count 
                FROM records 
                GROUP BY year 
                ORDER BY year
            ''')
            year_distribution = {}
            for row in cursor.fetchall():
                year_distribution[row['year']] = row['count']

            # Количество уникальных соавторов
            cursor.execute('SELECT COUNT(DISTINCT name) as count FROM coauthors')
            unique_coauthors = cursor.fetchone()['count']

            # Активность за последние 12 месяцев
            twelve_months_ago = (datetime.datetime.now() - datetime.timedelta(days=365)).strftime('%Y-%m-%d')
            cursor.execute('''
                SELECT strftime('%Y-%m', timestamp) as month, 
                       COUNT(*) as count 
                FROM activity_log 
                WHERE timestamp >= ?
                GROUP BY month 
                ORDER BY month
            ''', (twelve_months_ago,))
            monthly_activity = {}
            for row in cursor.fetchall():
                monthly_activity[row['month']] = row['count']

            # Последние 5 записей
            cursor.execute('''
                SELECT id, title, type, year, created_at 
                FROM records 
                ORDER BY created_at DESC 
                LIMIT 5
            ''')
            recent_records = []
            for row in cursor.fetchall():
                recent_records.append(dict(row))

            return {
                'type_distribution': type_distribution,
                'year_distribution': year_distribution,
                'unique_coauthors': unique_coauthors,
                'monthly_activity': monthly_activity,
                'recent_records': recent_records,
                'total_records': total_records
            }


class PortfolioApp:
    """Основной класс приложения с GUI"""

    def __init__(self, root):
        self.root = root
        self.root.title("Portfolio Manager")
        self.root.geometry("1200x700")

        # Инициализация базы данных
        self.db = PortfolioDatabase()

        # Текущая выбранная запись
        self.current_record_id = None

        # Для сортировки
        self.sort_direction = {}

        # Настройка стилей
        self.setup_styles()

        # Создание интерфейса
        self.create_widgets()

        # Загрузка записей
        self.load_records()

    def setup_styles(self):
        """Настраивает стили для виджетов"""
        style = ttk.Style()
        style.theme_use('clam')

        # Настройка цвета
        style.configure("Treeview",
                        background="#ffffff",
                        foreground="#333333",
                        rowheight=25,
                        fieldbackground="#ffffff")
        style.map('Treeview', background=[('selected', '#4a90e2')])

        style.configure("TButton", padding=6, relief="flat", background="#4a90e2")
        style.map("TButton", background=[('active', '#357ae8')])

    def create_widgets(self):
        """Создает все виджеты интерфейса"""
        # Основной контейнер
        main_container = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Левая панель (список записей)
        left_panel = ttk.Frame(main_container)
        main_container.add(left_panel, weight=1)

        # Правая панель (редактирование)
        right_panel = ttk.Frame(main_container)
        main_container.add(right_panel, weight=2)

        # 1. Панель управления (вверху слева)
        control_frame = ttk.LabelFrame(left_panel, text="Панель управления", padding=10)
        control_frame.pack(fill=tk.X, padx=5, pady=5)

        # Поля ввода
        ttk.Label(control_frame, text="Название:").grid(row=0, column=0, sticky=tk.W, pady=2)
        self.title_entry = ttk.Entry(control_frame, width=30)
        self.title_entry.grid(row=0, column=1, padx=5, pady=2)

        ttk.Label(control_frame, text="Тип записи:").grid(row=1, column=0, sticky=tk.W, pady=2)
        self.type_combobox = ttk.Combobox(control_frame,
                                          values=["Статья", "Книга", "Проект", "Доклад", "Патент", "Другое"],
                                          width=27)
        self.type_combobox.grid(row=1, column=1, padx=5, pady=2)
        self.type_combobox.set("Статья")

        ttk.Label(control_frame, text="Год:").grid(row=2, column=0, sticky=tk.W, pady=2)
        self.year_entry = ttk.Entry(control_frame, width=30)
        self.year_entry.grid(row=2, column=1, padx=5, pady=2)
        current_year = datetime.datetime.now().year
        self.year_entry.insert(0, str(current_year))

        # Кнопки управления
        button_frame = ttk.Frame(control_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=10)

        ttk.Button(button_frame, text="Создать",
                   command=self.create_record).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Сохранить",
                   command=self.save_record).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Удалить",
                   command=self.delete_record).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Открыть описание",
                   command=self.open_description).pack(side=tk.LEFT, padx=2)

        # Кнопки экспорта
        export_frame = ttk.Frame(control_frame)
        export_frame.grid(row=4, column=0, columnspan=2, pady=5)

        ttk.Button(export_frame, text="Экспорт в Excel",
                   command=self.export_to_excel).pack(side=tk.LEFT, padx=2)
        ttk.Button(export_frame, text="Экспорт в Word",
                   command=self.export_to_word).pack(side=tk.LEFT, padx=2)

        # 2. Список записей (Treeview)
        list_frame = ttk.LabelFrame(left_panel, text="Список записей", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Создаем Treeview с полосами прокрутки
        tree_scroll = ttk.Scrollbar(list_frame)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.records_tree = ttk.Treeview(list_frame,
                                         yscrollcommand=tree_scroll.set,
                                         selectmode="browse",
                                         columns=("title", "type", "year", "created_at"))
        tree_scroll.config(command=self.records_tree.yview)

        # Настраиваем колонки
        self.records_tree.heading("#0", text="ID")
        self.records_tree.column("#0", width=50, minwidth=50)

        columns = [("title", "Название", 200),
                   ("type", "Тип", 100),
                   ("year", "Год", 80),
                   ("created_at", "Дата создания", 120)]

        for col_id, heading, width in columns:
            self.records_tree.heading(col_id, text=heading,
                                      command=lambda c=col_id: self.sort_treeview(c))
            self.records_tree.column(col_id, width=width, minwidth=width)

        self.records_tree.pack(fill=tk.BOTH, expand=True)
        self.records_tree.bind('<<TreeviewSelect>>', self.on_record_select)

        # 3. Область редактирования (правая панель)
        edit_frame = ttk.LabelFrame(right_panel, text="Редактирование записи", padding=10)
        edit_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Подсказки по синтаксису
        help_frame = ttk.Frame(edit_frame)
        help_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(help_frame, text="Подсказки: ", font=("Arial", 9)).pack(side=tk.LEFT)
        ttk.Label(help_frame, text="Цитаты: > текст",
                  font=("Arial", 9, "italic"), foreground="#666").pack(side=tk.LEFT, padx=5)
        ttk.Label(help_frame, text="Код: ```код```",
                  font=("Arial", 9, "italic"), foreground="#666").pack(side=tk.LEFT, padx=5)
        ttk.Label(help_frame, text="Ссылки: [текст](url)",
                  font=("Arial", 9, "italic"), foreground="#666").pack(side=tk.LEFT, padx=5)

        # Текстовое поле для редактирования
        self.text_editor = scrolledtext.ScrolledText(edit_frame,
                                                     wrap=tk.WORD,
                                                     font=("Courier New", 10),
                                                     height=20)
        self.text_editor.pack(fill=tk.BOTH, expand=True)

        # 4. Панель соавторов
        coauthors_frame = ttk.LabelFrame(right_panel, text="Соавторы", padding=10)
        coauthors_frame.pack(fill=tk.X, padx=5, pady=5)

        # Поле для ввода соавтора
        input_frame = ttk.Frame(coauthors_frame)
        input_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(input_frame, text="Имя соавтора:").pack(side=tk.LEFT, padx=(0, 5))
        self.coauthor_entry = ttk.Entry(input_frame, width=40)
        self.coauthor_entry.pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(input_frame, text="Добавить соавтора",
                   command=self.add_coauthor).pack(side=tk.LEFT)

        # Область для отображения соавторов
        self.coauthors_label = ttk.Label(coauthors_frame,
                                         text="Соавторы не выбраны",
                                         wraplength=600,
                                         background="#f9f9f9",
                                         relief="solid",
                                         padding=5)
        self.coauthors_label.pack(fill=tk.X)

        # 5. Вкладка "Аналитика и отчётность"
        analytics_frame = ttk.LabelFrame(right_panel, text="Аналитика и отчётность", padding=10)
        analytics_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(analytics_frame, text="Сформировать отчёт",
                   command=self.generate_report,
                   width=20).pack(pady=10)

        # Статус бар
        self.status_bar = ttk.Label(self.root,
                                    text="Готово",
                                    relief=tk.SUNKEN,
                                    anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def load_records(self):
        """Загружает записи из БД в Treeview"""
        # Очищаем Treeview
        for item in self.records_tree.get_children():
            self.records_tree.delete(item)

        # Загружаем записи
        records = self.db.get_all_records()
        for record in records:
            self.records_tree.insert("", "end",
                                     iid=record['id'],
                                     text=str(record['id']),
                                     values=(record['title'],
                                             record['type'],
                                             record['year'],
                                             record['created_at']))

    def sort_treeview(self, col):
        """Сортирует Treeview по колонке"""
        items = [(self.records_tree.set(k, col), k) for k in self.records_tree.get_children('')]
        items.sort(reverse=self.sort_direction.get(col, False))

        for index, (val, k) in enumerate(items):
            self.records_tree.move(k, '', index)

        # Меняем направление сортировки
        self.sort_direction[col] = not self.sort_direction.get(col, False)

    def on_record_select(self, event):
        """Обрабатывает выбор записи в Treeview"""
        selection = self.records_tree.selection()
        if not selection:
            return

        self.current_record_id = int(selection[0])
        record = self.db.get_record(self.current_record_id)

        if record:
            # Заполняем поля ввода
            self.title_entry.delete(0, tk.END)
            self.title_entry.insert(0, record['title'])

            self.type_combobox.set(record['type'])

            self.year_entry.delete(0, tk.END)
            self.year_entry.insert(0, str(record['year']))

            # Загружаем содержимое файла
            file_path = record.get('abs_file_path')
            if file_path and os.path.exists(file_path):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    self.text_editor.delete(1.0, tk.END)
                    self.text_editor.insert(1.0, content)
                except Exception as e:
                    self.text_editor.delete(1.0, tk.END)
                    self.text_editor.insert(1.0, f"# Ошибка\n\nНе удалось загрузить файл: {str(e)}")
            else:
                self.text_editor.delete(1.0, tk.END)
                self.text_editor.insert(1.0, "# Ошибка\n\nФайл не найден")

            # Загружаем соавторов
            coauthors = self.db.get_coauthors(self.current_record_id)
            if coauthors:
                self.coauthors_label.config(text=", ".join(coauthors))
            else:
                self.coauthors_label.config(text="Соавторы отсутствуют")

            self.status_bar.config(text=f"Выбрана запись: {record['title']}")

    def create_record(self):
        """Создает новую запись"""
        title = self.title_entry.get().strip()
        record_type = self.type_combobox.get().strip()
        year = self.year_entry.get().strip()

        if not title:
            messagebox.showerror("Ошибка", "Введите название записи")
            return

        if not record_type:
            messagebox.showerror("Ошибка", "Выберите тип записи")
            return

        if not year.isdigit() or len(year) != 4:
            messagebox.showerror("Ошибка", "Введите корректный год (4 цифры)")
            return

        try:
            record_id = self.db.create_record(title, record_type, int(year))
            self.load_records()

            # Выделяем новую запись
            self.records_tree.selection_set(str(record_id))
            self.records_tree.see(str(record_id))

            self.status_bar.config(text=f"Создана запись: {title}")
            messagebox.showinfo("Успех", "Запись успешно создана")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать запись: {str(e)}")

    def save_record(self):
        """Сохраняет изменения в записи"""
        if not self.current_record_id:
            messagebox.showerror("Ошибка", "Выберите запись для сохранения")
            return

        title = self.title_entry.get().strip()
        record_type = self.type_combobox.get().strip()
        year = self.year_entry.get().strip()
        content = self.text_editor.get(1.0, tk.END).strip()

        if not title:
            messagebox.showerror("Ошибка", "Введите название записи")
            return

        if not record_type:
            messagebox.showerror("Ошибка", "Выберите тип записи")
            return

        if not year.isdigit() or len(year) != 4:
            messagebox.showerror("Ошибка", "Введите корректный год (4 цифры)")
            return

        try:
            success = self.db.update_record(self.current_record_id, title, record_type,
                                            int(year), content)
            if success:
                self.load_records()
                self.status_bar.config(text=f"Сохранена запись: {title}")
                messagebox.showinfo("Успех", "Запись успешно сохранена")
            else:
                messagebox.showerror("Ошибка", "Не удалось найти запись для обновления")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить запись: {str(e)}")

    def delete_record(self):
        """Удаляет запись"""
        if not self.current_record_id:
            messagebox.showerror("Ошибка", "Выберите запись для удаления")
            return

        # Подтверждение удаления
        if not messagebox.askyesno("Подтверждение", "Вы уверены, что хотите удалить эту запись?"):
            return

        try:
            success = self.db.delete_record(self.current_record_id)
            if success:
                self.load_records()
                self.current_record_id = None

                # Очищаем поля
                self.title_entry.delete(0, tk.END)
                self.type_combobox.set("Статья")
                self.year_entry.delete(0, tk.END)
                self.year_entry.insert(0, str(datetime.datetime.now().year))
                self.text_editor.delete(1.0, tk.END)
                self.coauthors_label.config(text="Соавторы не выбраны")

                self.status_bar.config(text="Запись удалена")
                messagebox.showinfo("Успех", "Запись успешно удалена")
            else:
                messagebox.showerror("Ошибка", "Не удалось найти запись для удаления")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить запись: {str(e)}")

    def open_description(self):
        """Открывает описание во внешнем просмотрщике"""
        if not self.current_record_id:
            messagebox.showerror("Ошибка", "Выберите запись")
            return

        record = self.db.get_record(self.current_record_id)

        file_path = record.get('abs_file_path') if record else None

        if file_path and os.path.exists(file_path):
            try:
                # Преобразуем markdown в HTML
                with open(file_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()

                html_content = markdown.markdown(md_content)

                # Создаем временный HTML файл
                temp_file = "temp_preview.html"
                with open(temp_file, 'w', encoding='utf-8') as f:
                    f.write(f"""
                    <html>
                    <head>
                        <meta charset="utf-8">
                        <title>{record['title']}</title>
                        <style>
                            body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                            h1 {{ color: #333; }}
                            pre {{ background: #f4f4f4; padding: 10px; border-radius: 5px; }}
                            blockquote {{ border-left: 3px solid #ccc; padding-left: 10px; color: #666; }}
                        </style>
                    </head>
                    <body>
                    {html_content}
                    </body>
                    </html>
                    """)

                # Открываем в браузере
                webbrowser.open(f"file://{os.path.abspath(temp_file)}")

            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось открыть описание: {str(e)}")
        else:
            messagebox.showerror("Ошибка", "Файл записи не найден")

    def add_coauthor(self):
        """Добавляет соавтора к текущей записи"""
        if not self.current_record_id:
            messagebox.showerror("Ошибка", "Выберите запись")
            return

        name = self.coauthor_entry.get().strip()
        if not name:
            messagebox.showerror("Ошибка", "Введите имя соавтора")
            return

        try:
            success = self.db.add_coauthor(self.current_record_id, name)
            if success:
                # Обновляем список соавторов
                coauthors = self.db.get_coauthors(self.current_record_id)
                self.coauthors_label.config(text=", ".join(coauthors))
                self.coauthor_entry.delete(0, tk.END)
                self.status_bar.config(text=f"Добавлен соавтор: {name}")
            else:
                messagebox.showwarning("Предупреждение", "Этот соавтор уже добавлен")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить соавтора: {str(e)}")

    def generate_report(self):
        """Генерирует отчет в Excel и Word"""
        try:
            # Создаем директорию для отчетов
            os.makedirs('reports', exist_ok=True)

            self.status_bar.config(text="Сбор данных для отчета...")
            self.root.update()

            # Получаем статистику
            stats = self.db.get_statistics()

            if stats['total_records'] == 0:
                messagebox.showwarning("Предупреждение", "Нет данных для формирования отчета")
                self.status_bar.config(text="Нет данных для отчета")
                return

            self.status_bar.config(text="Генерация Excel отчета...")
            self.root.update()

            # Генерируем Excel отчет
            self.generate_excel_report(stats)

            self.status_bar.config(text="Генерация Word отчета...")
            self.root.update()

            # Генерируем Word отчет
            self.generate_word_report(stats)

            self.status_bar.config(text="Отчеты успешно сгенерированы")
            messagebox.showinfo("Успех",
                                f"Отчеты успешно сгенерированы в папке 'reports'\n"
                                f"Всего записей: {stats['total_records']}")

        except Exception as e:
            error_msg = f"Не удалось сгенерировать отчет: {str(e)}"
            messagebox.showerror("Ошибка", error_msg)
            self.status_bar.config(text=error_msg)

    def generate_excel_report(self, stats):
        """Генерирует Excel отчет с графиками"""
        # Создаем новую рабочую книгу
        wb = Workbook()

        # Лист "Статистика"
        ws_stats = wb.active
        ws_stats.title = "Статистика"

        # Заголовок
        ws_stats.append(["Отчет по портфолио", ""])
        ws_stats.append([f"Дата формирования: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}", ""])
        ws_stats.append([])

        # Ключевые показатели
        ws_stats.append(["Ключевые показатели"])
        ws_stats.append(["Всего записей:", stats['total_records']])
        ws_stats.append(["Уникальных соавторов:", stats['unique_coauthors']])

        # Период записей
        if stats['year_distribution']:
            years = list(stats['year_distribution'].keys())
            min_year = min(years)
            max_year = max(years)
            ws_stats.append(["Период записей:", f"{min_year} - {max_year}"])
        ws_stats.append([])

        # Распределение по типам
        ws_stats.append(["Распределение записей по типам"])
        ws_stats.append(["Тип", "Количество"])
        for record_type, count in stats['type_distribution'].items():
            ws_stats.append([record_type, count])
        ws_stats.append([])

        # Распределение по годам
        ws_stats.append(["Распределение записей по годам"])
        ws_stats.append(["Год", "Количество"])
        for year, count in stats['year_distribution'].items():
            ws_stats.append([year, count])
        ws_stats.append([])

        # Активность по месяцам
        ws_stats.append(["Активность за последние 12 месяцев"])
        ws_stats.append(["Месяц", "Количество действий"])
        for month, count in stats['monthly_activity'].items():
            ws_stats.append([month, count])

        # Лист "Графики"
        ws_charts = wb.create_sheet(title="Графики")

        # Создаем графики
        self.create_charts(stats, ws_charts)

        # Сохраняем файл
        excel_path = os.path.join('reports', 'portfolio_report.xlsx')
        wb.save(excel_path)

    def create_charts(self, stats, ws_charts):
        """Создает и сохраняет графики"""
        # Создаем папку для графиков
        charts_dir = os.path.join('reports', 'charts')
        os.makedirs(charts_dir, exist_ok=True)

        # 1. График распределения по типам
        if stats['type_distribution']:
            plt.figure(figsize=(10, 6))
            types = list(stats['type_distribution'].keys())
            counts = list(stats['type_distribution'].values())

            bars = plt.bar(types, counts)
            plt.title('Распределение записей по типам')
            plt.xlabel('Тип записи')
            plt.ylabel('Количество')
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()

            # Добавляем значения на столбцы
            for bar, count in zip(bars, counts):
                plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height(),
                         str(count), ha='center', va='bottom')

            # Сохраняем график
            chart1_path = os.path.join(charts_dir, 'chart_types.png')
            plt.savefig(chart1_path, dpi=100, bbox_inches='tight')
            plt.close()

            # Добавляем в Excel
            img = XLImage(chart1_path)
            ws_charts.add_image(img, 'A1')

        # 2. График распределения по годам
        if stats['year_distribution']:
            plt.figure(figsize=(10, 6))
            years = [str(y) for y in sorted(stats['year_distribution'].keys())]
            counts = [stats['year_distribution'][int(y)] for y in years]

            bars = plt.bar(years, counts)
            plt.title('Распределение записей по годам')
            plt.xlabel('Год')
            plt.ylabel('Количество')
            plt.tight_layout()

            # Добавляем значения на столбцы
            for bar, count in zip(bars, counts):
                plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height(),
                         str(count), ha='center', va='bottom')

            # Сохраняем график
            chart2_path = os.path.join(charts_dir, 'chart_years.png')
            plt.savefig(chart2_path, dpi=100, bbox_inches='tight')
            plt.close()

            # Добавляем в Excel
            img = XLImage(chart2_path)
            ws_charts.add_image(img, 'A30')

    def generate_word_report(self, stats):
        """Генерирует Word отчет"""
        # Создаем новый документ
        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5

        # Титульный лист
        title = doc.add_heading('Отчет по портфолио', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(f'Дата формирования: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Сводная таблица ключевых показателей
        doc.add_heading('Ключевые показатели', level=1)

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Shading'

        table.cell(0, 0).text = 'Показатель'
        table.cell(0, 1).text = 'Значение'
        table.cell(1, 0).text = 'Всего записей'
        table.cell(1, 1).text = str(stats['total_records'])
        table.cell(2, 0).text = 'Уникальных соавторов'
        table.cell(2, 1).text = str(stats['unique_coauthors'])

        # Период записей
        if stats['year_distribution']:
            years = list(stats['year_distribution'].keys())
            min_year = min(years)
            max_year = max(years)
            table.cell(3, 0).text = 'Период записей'
            table.cell(3, 1).text = f"{min_year} - {max_year}"
        else:
            table.cell(3, 0).text = 'Период записей'
            table.cell(3, 1).text = 'Нет данных'

        # Распределение по типам
        doc.add_heading('Распределение записей по типам', level=2)
        if stats['type_distribution']:
            table_type = doc.add_table(rows=len(stats['type_distribution']) + 1, cols=2)
            table_type.style = 'Light List'

            table_type.cell(0, 0).text = 'Тип записи'
            table_type.cell(0, 1).text = 'Количество'

            for i, (record_type, count) in enumerate(stats['type_distribution'].items(), 1):
                table_type.cell(i, 0).text = record_type
                table_type.cell(i, 1).text = str(count)

        # График
        doc.add_heading('Визуализация данных', level=2)
        doc.add_paragraph('График распределения записей по годам:')

        chart_path = os.path.join('reports', 'charts', 'chart_years.png')
        if os.path.exists(chart_path):
            try:
                doc.add_picture(chart_path, width=Inches(6))
            except:
                doc.add_paragraph("(График не удалось вставить)")
        else:
            doc.add_paragraph("График не доступен")

        # ПОЛНЫЕ ЗАПИСИ
        doc.add_page_break()
        doc.add_heading('Полный список записей', level=2)

        # Получаем все записи
        all_records = self.db.get_all_records()

        if all_records:
            for record in all_records:
                doc.add_heading(record['title'], level=3)

                # Добавляем метаданные
                meta_text = f"Тип: {record['type']} | Год: {record['year']} | Создано: {record['created_at']}"

                # Соавторы
                coauthors = record.get('coauthors')
                if coauthors:
                    meta_text += f" | Соавторы: {coauthors}"
                else:
                    meta_text += " | Соавторы: Отсутствуют"

                meta = doc.add_paragraph(meta_text)

                # Получаем содержимое файла
                file_path = record.get('abs_file_path')
                if file_path and os.path.exists(file_path):
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()

                        # Добавляем содержимое как обычный текст
                        if content.strip():
                            doc.add_heading('Описание:', level=4)

                            # Разбиваем на абзацы и добавляем
                            paragraphs = content.split('\n\n')
                            for para in paragraphs:
                                if para.strip():
                                    # Простая обработка маркдауна
                                    lines = para.split('\n')
                                    for line in lines:
                                        line = line.strip()
                                        if line.startswith('# '):
                                            doc.add_heading(line[2:].strip(), level=1)
                                        elif line.startswith('## '):
                                            doc.add_heading(line[3:].strip(), level=2)
                                        elif line.startswith('### '):
                                            doc.add_heading(line[4:].strip(), level=3)
                                        elif line.startswith('> '):
                                            quote = doc.add_paragraph(style='Intense Quote')
                                            quote.add_run(line[2:].strip())
                                        elif line.startswith('```') or line.endswith('```'):
                                            # Пропускаем блоки кода в Word отчете
                                            continue
                                        elif line:
                                            doc.add_paragraph(line)
                    except Exception as e:
                        doc.add_paragraph(f"Ошибка при загрузке описания: {str(e)}")
                else:
                    doc.add_paragraph("Файл записи не найден")

                # Разделитель между записями
                doc.add_paragraph()
                doc.add_paragraph("—" * 50)
                doc.add_paragraph()
        else:
            doc.add_paragraph("Нет записей для отображения")

        # Сохраняем документ
        word_path = os.path.join('reports', 'portfolio_report.docx')
        doc.save(word_path)

    def export_to_excel(self):
        """Экспортирует все записи в Excel"""
        try:
            records = self.db.get_all_records()
            if not records:
                messagebox.showwarning("Предупреждение", "Нет записей для экспорта")
                return

            # Создаем новую рабочую книгу
            wb = Workbook()
            ws = wb.active
            ws.title = "Портфолио"

            # Заголовки
            headers = ['ID', 'Название', 'Тип', 'Год', 'Дата создания', 'Соавторы']
            ws.append(headers)

            # Данные
            for record in records:
                ws.append([
                    record['id'],
                    record['title'],
                    record['type'],
                    record['year'],
                    record['created_at'],
                    record.get('coauthors', '') or ''
                ])

            # Настраиваем ширину колонок
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            # Сохраняем файл
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                initialfile="portfolio_export.xlsx"
            )

            if file_path:
                wb.save(file_path)
                self.status_bar.config(text=f"Данные экспортированы в {file_path}")
                messagebox.showinfo("Успех", "Данные успешно экспортированы в Excel")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать в Excel: {str(e)}")

    def export_to_word(self):
        """Экспортирует все записи в Word"""
        try:
            records = self.db.get_all_records()
            if not records:
                messagebox.showwarning("Предупреждение", "Нет записей для экспорта")
                return

            # Создаем новый документ
            doc = Document()

            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.5

            # Титульная страница
            title = doc.add_heading('Экспорт портфолио', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f'Дата экспорта: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}')
            doc.add_paragraph(f'Всего записей: {len(records)}')
            doc.add_paragraph()

            # Добавляем записи
            for record in records:
                # Заголовок записи
                doc.add_heading(record['title'], level=2)

                # Таблица с метаданными
                table = doc.add_table(rows=4, cols=2)
                table.style = 'Light Grid Accent 1'

                table.cell(0, 0).text = 'Тип:'
                table.cell(0, 1).text = record['type']
                table.cell(1, 0).text = 'Год:'
                table.cell(1, 1).text = str(record['year'])
                table.cell(2, 0).text = 'Дата создания:'
                table.cell(2, 1).text = record['created_at']
                table.cell(3, 0).text = 'Соавторы:'
                table.cell(3, 1).text = record.get('coauthors', '') or 'Отсутствуют'

                # Добавляем содержимое записи
                doc.add_paragraph()
                doc.add_heading('Содержание:', level=3)

                # Получаем абсолютный путь к файлу
                file_path = record.get('abs_file_path')

                # Читаем содержимое файла
                if file_path and os.path.exists(file_path):
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()

                        # Добавляем содержимое как обычный текст
                        paragraphs = content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                lines = para.split('\n')
                                for line in lines:
                                    line = line.strip()
                                    if line and not line.startswith('```'):
                                        doc.add_paragraph(line)
                    except Exception as e:
                        doc.add_paragraph(f"Ошибка при чтении файла: {str(e)}")
                else:
                    doc.add_paragraph("Файл записи не найден")

                # Разделитель между записями
                doc.add_page_break()

            # Сохраняем файл
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
                initialfile="portfolio_export.docx"
            )

            if file_path:
                doc.save(file_path)
                self.status_bar.config(text=f"Данные экспортированы в {file_path}")
                messagebox.showinfo("Успех", "Данные успешно экспортированы в Word")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать в Word: {str(e)}")


def main():
    """Основная функция запуска приложения"""
    # Создаем главное окно
    root = tk.Tk()

    # Инициализируем приложение
    app = PortfolioApp(root)

    # Запускаем главный цикл
    root.mainloop()


if __name__ == "__main__":
    main()