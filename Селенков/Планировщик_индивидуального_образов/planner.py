import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
from docx.shared import Pt, Inches, RGBColor  # Добавьте RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

# Проверяем, установлен ли PostgreSQL
try:
    import psycopg2
    from psycopg2 import sql

    POSTGRES_AVAILABLE = True
except ImportError:
    POSTGRES_AVAILABLE = False


class EducationalPlanner:
    def __init__(self, root):
        self.root = root
        self.root.title("Планировщик индивидуального образовательного маршрута")
        self.root.geometry("1200x700")

        # Настройки по умолчанию
        self.specialty = None
        self.current_goals = []
        self.skills_autocomplete = []

        # Определяем тип БД
        self.db_type = self.detect_db_type()

        # Подключаемся к БД
        self.conn = self.connect_to_db()
        self.cursor = self.conn.cursor()

        # Инициализируем БД
        self.init_database()

        # Загружаем данные
        self.load_skills_autocomplete()

        # Создаем интерфейс
        self.create_widgets()

        # Проверяем достижения
        self.check_achievements()

    def detect_db_type(self):
        """Определяем, какую БД использовать"""
        if POSTGRES_AVAILABLE:
            try:
                # Пытаемся подключиться к PostgreSQL
                conn = psycopg2.connect(
                    host='localhost',
                    database='postgres',
                    user='postgres',
                    password='1111',
                    port='5432'
                )
                conn.close()
                print("Используется PostgreSQL")
                return 'postgres'
            except Exception as e:
                print(f"PostgreSQL недоступен: {e}. Используется SQLite")
                return 'sqlite'
        print("Используется SQLite")
        return 'sqlite'

    def connect_to_db(self):
        """Подключение к базе данных"""
        if self.db_type == 'postgres':
            return psycopg2.connect(
                host='localhost',
                database='postgres',
                user='postgres',
                password='1111',
                port='5432'
            )
        else:
            return sqlite3.connect('educational_planner.db')

    def init_database(self):
        """Инициализация базы данных"""
        # Таблица цели
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS цели (
                id SERIAL PRIMARY KEY,
                название TEXT NOT NULL,
                тип TEXT NOT NULL,
                статус TEXT NOT NULL,
                план_дата TEXT,
                факт_дата TEXT,
                описание TEXT
            )
        ''' if self.db_type == 'postgres' else '''
            CREATE TABLE IF NOT EXISTS цели (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                название TEXT NOT NULL,
                тип TEXT NOT NULL,
                статус TEXT NOT NULL,
                план_дата TEXT,
                факт_дата TEXT,
                описание TEXT
            )
        ''')

        # Таблица навыки
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS навыки (
                id SERIAL PRIMARY KEY,
                название TEXT UNIQUE NOT NULL
            )
        ''' if self.db_type == 'postgres' else '''
            CREATE TABLE IF NOT EXISTS навыки (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                название TEXT UNIQUE NOT NULL
            )
        ''')

        # Таблица цель_навыки
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS цель_навыки (
                id SERIAL PRIMARY KEY,
                цель_id INTEGER,
                навык_id INTEGER
            )
        ''' if self.db_type == 'postgres' else '''
            CREATE TABLE IF NOT EXISTS цель_навыки (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                цель_id INTEGER,
                навык_id INTEGER
            )
        ''')

        # Таблица компетенции
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS компетенции (
                id SERIAL PRIMARY KEY,
                название TEXT NOT NULL,
                категория TEXT
            )
        ''' if self.db_type == 'postgres' else '''
            CREATE TABLE IF NOT EXISTS компетенции (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                название TEXT NOT NULL,
                категория TEXT
            )
        ''')

        # Таблица цель_компетенции
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS цель_компетенции (
                id SERIAL PRIMARY KEY,
                цель_id INTEGER,
                компетенция_id INTEGER,
                уровень INTEGER CHECK (уровень >= 0 AND уровень <= 5)
            )
        ''' if self.db_type == 'postgres' else '''
            CREATE TABLE IF NOT EXISTS цель_компетенции (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                цель_id INTEGER,
                компетенция_id INTEGER,
                уровень INTEGER CHECK (уровень >= 0 AND уровень <= 5)
            )
        ''')

        # Таблица достижения
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS достижения (
                код TEXT PRIMARY KEY,
                название TEXT NOT NULL,
                описание TEXT,
                получено INTEGER DEFAULT 0
            )
        ''')

        # Таблица цель_на_семестр
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS цель_на_семестр (
                id SERIAL PRIMARY KEY,
                текст_цели TEXT NOT NULL,
                тип_цели TEXT,
                параметр TEXT,
                текущий_прогресс INTEGER DEFAULT 0,
                целевой_прогресс INTEGER NOT NULL
            )
        ''' if self.db_type == 'postgres' else '''
            CREATE TABLE IF NOT EXISTS цель_на_семестр (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                текст_цели TEXT NOT NULL,
                тип_цели TEXT,
                параметр TEXT,
                текущий_прогресс INTEGER DEFAULT 0,
                целевой_прогресс INTEGER NOT NULL
            )
        ''')

        self.conn.commit()

        # Заполняем таблицу достижений начальными данными
        achievements = [
            ('start', 'Старт', 'Создана хотя бы одна цель', 0),
            ('punctual', 'Пунктуальный', 'Три или более завершённых цели с фактической датой не позже плановой', 0),
            ('multitarget', 'Многоцелевой', 'Есть цели минимум трёх разных типов', 0),
            ('skill_growth', 'Навыковый рост', 'У одного навыка четыре или более связанных завершённых целей', 0),
            ('planner', 'Планировщик', 'Одновременно в статусе "В процессе" пять или более целей', 0)
        ]

        for achievement in achievements:
            # Проверяем, существует ли уже достижение
            self.cursor.execute("SELECT COUNT(*) FROM достижения WHERE код = %s"
                                if self.db_type == 'postgres' else
                                "SELECT COUNT(*) FROM достижения WHERE код = ?",
                                (achievement[0],))

            if self.cursor.fetchone()[0] == 0:
                # Вставляем только если не существует
                self.cursor.execute('''
                    INSERT INTO достижения (код, название, описание, получено)
                    VALUES (%s, %s, %s, %s)
                ''' if self.db_type == 'postgres' else '''
                    INSERT INTO достижения (код, название, описание, получено)
                    VALUES (?, ?, ?, ?)
                ''', achievement)

        # Загружаем компетенции из JSON файла, если таблица пуста
        self.cursor.execute("SELECT COUNT(*) FROM компетенции")
        if self.cursor.fetchone()[0] == 0:
            self.load_competencies_from_json()

        self.conn.commit()

    def load_competencies_from_json(self):
        """Загрузка компетенций из JSON файла"""
        try:
            if os.path.exists('competencies.json'):
                with open('competencies.json', 'r', encoding='utf-8') as f:
                    competencies = json.load(f)
                    for comp in competencies:
                        self.cursor.execute(
                            "INSERT INTO компетенции (название, категория) VALUES (%s, %s)"
                            if self.db_type == 'postgres' else
                            "INSERT INTO компетенции (название, категория) VALUES (?, ?)",
                            (comp['название'], comp.get('категория', ''))
                        )
                self.conn.commit()
                print("Компетенции загружены из JSON файла")
        except Exception as e:
            print(f"Ошибка загрузки компетенций: {e}")

    def load_skills_autocomplete(self):
        """Загрузка списка навыков для автодополнения"""
        self.cursor.execute("SELECT название FROM навыки")
        self.skills_autocomplete = [row[0] for row in self.cursor.fetchall()]

    def create_widgets(self):
        """Создание элементов интерфейса"""
        # Создаем Notebook (вкладки)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Вкладка "Мои цели"
        self.tab_goals = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_goals, text='Мои цели')
        self.create_goals_tab()

        # Вкладка "Мой профиль"
        self.tab_profile = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_profile, text='Мой профиль')
        self.create_profile_tab()

        # Вкладка "Компетенции"
        self.tab_competencies = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_competencies, text='Компетенции')
        self.create_competencies_tab()

        # Вкладка "Достижения"
        self.tab_achievements = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_achievements, text='Достижения')
        self.create_achievements_tab()

        # Вкладка "Цели на семестр"
        self.tab_semester = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_semester, text='Цели на семестр')
        self.create_semester_tab()

        # Вкладка "Настройки"
        self.tab_settings = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_settings, text='Настройки')
        self.create_settings_tab()

    def create_goals_tab(self):
        """Создание вкладки Мои цели"""
        # Панель управления
        control_frame = ttk.Frame(self.tab_goals)
        control_frame.pack(fill='x', padx=10, pady=5)

        ttk.Button(control_frame, text="Добавить цель",
                   command=self.add_goal).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Редактировать",
                   command=self.edit_goal).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Удалить",
                   command=self.delete_goal).pack(side='left', padx=5)

        # Дерево целей
        columns = ('ID', 'Название', 'Тип', 'Статус', 'План дата', 'Факт дата')
        self.goals_tree = ttk.Treeview(self.tab_goals, columns=columns, show='headings')

        for col in columns:
            self.goals_tree.heading(col, text=col)
            self.goals_tree.column(col, width=150)

        scrollbar = ttk.Scrollbar(self.tab_goals, orient='vertical',
                                  command=self.goals_tree.yview)
        self.goals_tree.configure(yscrollcommand=scrollbar.set)

        self.goals_tree.pack(side='left', fill='both', expand=True, padx=(10, 0), pady=10)
        scrollbar.pack(side='right', fill='y', padx=(0, 10), pady=10)

        self.load_goals()

    def create_profile_tab(self):
        """Создание вкладки Мой профиль"""
        # Фрейм для статистики
        stats_frame = ttk.LabelFrame(self.tab_profile, text="Статистика")
        stats_frame.pack(fill='both', expand=True, padx=10, pady=10)

        # Навыки
        skills_frame = ttk.LabelFrame(stats_frame, text="Навыки")
        skills_frame.pack(fill='both', padx=5, pady=5)

        self.skills_text = tk.Text(skills_frame, height=10, width=50)
        self.skills_text.pack(padx=5, pady=5)

        # Статистика по типам целей
        types_frame = ttk.LabelFrame(stats_frame, text="Статистика по типам целей")
        types_frame.pack(fill='both', padx=5, pady=5)

        self.types_text = tk.Text(types_frame, height=5, width=50)
        self.types_text.pack(padx=5, pady=5)

        # Процент целей завершённых в срок
        timely_frame = ttk.LabelFrame(stats_frame, text="Своевременность")
        timely_frame.pack(fill='both', padx=5, pady=5)

        self.timely_label = ttk.Label(timely_frame, text="")
        self.timely_label.pack(padx=5, pady=5)

        # Кнопка обновления
        ttk.Button(stats_frame, text="Обновить статистику",
                   command=self.update_profile).pack(pady=10)

    def create_competencies_tab(self):
        """Создание вкладки Компетенции"""
        # Средний уровень по компетенциям
        avg_frame = ttk.LabelFrame(self.tab_competencies, text="Средний уровень по компетенциям")
        avg_frame.pack(fill='both', padx=10, pady=5)

        self.avg_text = tk.Text(avg_frame, height=15, width=60)
        scrollbar = ttk.Scrollbar(avg_frame, command=self.avg_text.yview)
        self.avg_text.configure(yscrollcommand=scrollbar.set)

        self.avg_text.pack(side='left', fill='both', expand=True, padx=5, pady=5)
        scrollbar.pack(side='right', fill='y', padx=(0, 5), pady=5)

        # Слабые зоны
        weak_frame = ttk.LabelFrame(self.tab_competencies, text="Слабые зоны")
        weak_frame.pack(fill='both', padx=10, pady=5)

        self.weak_text = tk.Text(weak_frame, height=5, width=60)
        self.weak_text.pack(padx=5, pady=5)

        # Рекомендации
        rec_frame = ttk.LabelFrame(self.tab_competencies, text="Рекомендации")
        rec_frame.pack(fill='both', padx=10, pady=5)

        self.rec_text = tk.Text(rec_frame, height=5, width=60)
        self.rec_text.pack(padx=5, pady=5)

        # Кнопка обновления
        ttk.Button(self.tab_competencies, text="Обновить",
                   command=self.update_competencies).pack(pady=10)

    def create_achievements_tab(self):
        """Создание вкладки Достижения"""
        frame = ttk.Frame(self.tab_achievements)
        frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.achievements_text = tk.Text(frame, height=20, width=70)
        scrollbar = ttk.Scrollbar(frame, command=self.achievements_text.yview)
        self.achievements_text.configure(yscrollcommand=scrollbar.set)

        self.achievements_text.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        self.update_achievements()

    def create_semester_tab(self):
        """Создание вкладки Цели на семестр"""
        # Панель управления
        control_frame = ttk.Frame(self.tab_semester)
        control_frame.pack(fill='x', padx=10, pady=5)

        ttk.Button(control_frame, text="Добавить цель",
                   command=self.add_semester_goal).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Обновить прогресс",
                   command=self.update_semester_progress).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Сформировать отчёт",
                   command=self.generate_report).pack(side='right', padx=5)

        # Список целей на семестр
        self.semester_tree = ttk.Treeview(self.tab_semester,
                                          columns=('ID', 'Цель', 'Тип', 'Прогресс', 'Целевой'),
                                          show='headings')

        for col in ('ID', 'Цель', 'Тип', 'Прогресс', 'Целевой'):
            self.semester_tree.heading(col, text=col)

        scrollbar = ttk.Scrollbar(self.tab_semester, command=self.semester_tree.yview)
        self.semester_tree.configure(yscrollcommand=scrollbar.set)

        self.semester_tree.pack(side='left', fill='both', expand=True, padx=(10, 0), pady=10)
        scrollbar.pack(side='right', fill='y', padx=(0, 10), pady=10)

        self.load_semester_goals()

    def create_settings_tab(self):
        """Создание вкладки Настройки"""
        frame = ttk.Frame(self.tab_settings)
        frame.pack(fill='both', expand=True, padx=50, pady=50)

        # Выбор специальности
        ttk.Label(frame, text="Выберите специальность:").pack(pady=10)

        self.specialty_var = tk.StringVar()
        specialties = ["Информационные системы", "Программная инженерия",
                       "Информационная безопасность", "Прикладная информатика"]

        specialty_combo = ttk.Combobox(frame, textvariable=self.specialty_var,
                                       values=specialties, state='readonly')
        specialty_combo.pack(pady=10)

        ttk.Button(frame, text="Сохранить",
                   command=self.save_specialty).pack(pady=20)

        # Информация о БД
        db_info = f"Используемая БД: {self.db_type.upper()}"
        ttk.Label(frame, text=db_info).pack(pady=10)

    def load_goals(self):
        """Загрузка целей в дерево"""
        # Очищаем дерево
        for item in self.goals_tree.get_children():
            self.goals_tree.delete(item)

        # Загружаем данные
        self.cursor.execute('''
            SELECT id, название, тип, статус, план_дата, факт_дата 
            FROM цели ORDER BY план_дата
        ''')

        for row in self.cursor.fetchall():
            self.goals_tree.insert('', 'end', values=row)

    def add_goal(self):
        """Добавление новой цели"""
        dialog = GoalDialog(self.root, "Добавить цель", self, None)
        self.root.wait_window(dialog.dialog)
        self.load_goals()
        self.check_achievements()

    def edit_goal(self):
        """Редактирование выбранной цели"""
        selection = self.goals_tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите цель для редактирования")
            return

        item = self.goals_tree.item(selection[0])
        goal_id = item['values'][0]

        dialog = GoalDialog(self.root, "Редактировать цель", self, goal_id)
        self.root.wait_window(dialog.dialog)
        self.load_goals()
        self.check_achievements()

    def delete_goal(self):
        """Удаление выбранной цели"""
        selection = self.goals_tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите цель для удаления")
            return

        if messagebox.askyesno("Подтверждение", "Удалить выбранную цель?"):
            item = self.goals_tree.item(selection[0])
            goal_id = item['values'][0]

            # Удаляем связи
            self.cursor.execute("DELETE FROM цель_навыки WHERE цель_id = %s"
                                if self.db_type == 'postgres' else
                                "DELETE FROM цель_навыки WHERE цель_id = ?",
                                (goal_id,))
            self.cursor.execute("DELETE FROM цель_компетенции WHERE цель_id = %s"
                                if self.db_type == 'postgres' else
                                "DELETE FROM цель_компетенции WHERE цель_id = ?",
                                (goal_id,))
            self.cursor.execute("DELETE FROM цели WHERE id = %s"
                                if self.db_type == 'postgres' else
                                "DELETE FROM цели WHERE id = ?",
                                (goal_id,))
            self.conn.commit()

            self.load_goals()
            self.check_achievements()
            messagebox.showinfo("Успех", "Цель удалена")

    def update_profile(self):
        """Обновление статистики профиля"""
        # Очищаем поля
        self.skills_text.delete(1.0, tk.END)
        self.types_text.delete(1.0, tk.END)

        # Статистика по навыкам
        self.cursor.execute('''
            SELECT н.название, COUNT(цн.цель_id) as количество
            FROM навыки н
            LEFT JOIN цель_навыки цн ON н.id = цн.навык_id
            LEFT JOIN цели ц ON цн.цель_id = ц.id AND ц.статус = 'Завершено'
            GROUP BY н.id
            ORDER BY количество DESC
        ''')

        skills_stats = self.cursor.fetchall()
        for skill, count in skills_stats:
            self.skills_text.insert(tk.END, f"{skill} - {count} целей\n")

        # Статистика по типам целей
        self.cursor.execute('''
            SELECT тип, 
                   COUNT(CASE WHEN статус = 'Завершено' THEN 1 END) as завершено,
                   COUNT(*) as всего
            FROM цели
            GROUP BY тип
        ''')

        types_stats = self.cursor.fetchall()
        for type_name, completed, total in types_stats:
            self.types_text.insert(tk.END, f"{type_name}: {completed} из {total} завершено\n")

        # Процент целей завершённых в срок
        self.cursor.execute('''
            SELECT 
                COUNT(CASE WHEN факт_дата <= план_дата AND статус = 'Завершено' THEN 1 END) as в_срок,
                COUNT(CASE WHEN статус = 'Завершено' THEN 1 END) as всего_завершено
            FROM цели
            WHERE план_дата IS NOT NULL AND факт_дата IS NOT NULL
        ''')

        timely, total_completed = self.cursor.fetchone()
        if total_completed and total_completed > 0:
            percentage = (timely / total_completed) * 100
            self.timely_label.config(text=f"Завершено в срок: {percentage:.1f}%")
        else:
            self.timely_label.config(text="Завершено в срок: нет данных")

    def update_competencies(self):
        """Обновление информации о компетенциях"""
        # Очищаем поля
        self.avg_text.delete(1.0, tk.END)
        self.weak_text.delete(1.0, tk.END)
        self.rec_text.delete(1.0, tk.END)

        # Средний уровень по компетенциям
        self.cursor.execute('''
            SELECT к.название, к.категория, AVG(цк.уровень) as средний_уровень
            FROM компетенции к
            LEFT JOIN цель_компетенции цк ON к.id = цк.компетенция_id
            LEFT JOIN цели ц ON цк.цель_id = ц.id
            WHERE ц.статус = 'Завершено' OR ц.статус IS NULL
            GROUP BY к.id, к.название, к.категория
            ORDER BY средний_уровень DESC
        ''')

        competencies = self.cursor.fetchall()
        for name, category, avg_level in competencies:
            if avg_level is None:
                avg_level = 0
            self.avg_text.insert(tk.END,
                                 f"{name} ({category}): {avg_level:.1f}\n")

            # Слабые зоны
            if avg_level < 3:
                self.weak_text.insert(tk.END, f"• {name}: {avg_level:.1f}\n")

            # Рекомендации
            if avg_level < 2:
                self.rec_text.insert(tk.END,
                                     f"Вы почти не развиваете компетенцию {name}. "
                                     f"Рекомендуем добавить цели для её развития.\n\n")
            elif avg_level < 3:
                self.rec_text.insert(tk.END,
                                     f"Компетенция {name} требует внимания. "
                                     f"Рекомендуем практические задания.\n\n")

    def update_achievements(self):
        """Обновление списка достижений"""
        self.achievements_text.delete(1.0, tk.END)

        self.cursor.execute('''
            SELECT название, описание, получено FROM достижения
            ORDER BY получено DESC, код
        ''')

        achievements = self.cursor.fetchall()
        for name, desc, obtained in achievements:
            status = "✓ ПОЛУЧЕНО" if obtained else "○ не получено"
            self.achievements_text.insert(tk.END,
                                          f"{status}\n{name}\n{desc}\n\n")

    def load_semester_goals(self):
        """Загрузка целей на семестр"""
        for item in self.semester_tree.get_children():
            self.semester_tree.delete(item)

        self.cursor.execute('''
            SELECT id, текст_цели, тип_цели, текущий_прогресс, целевой_прогресс
            FROM цель_на_семестр
        ''')

        for row in self.cursor.fetchall():
            progress = f"{row[3]} из {row[4]}"
            self.semester_tree.insert('', 'end', values=(row[0], row[1], row[2], progress, row[4]))

    def add_semester_goal(self):
        """Добавление цели на семестр"""
        dialog = SemesterGoalDialog(self.root, self)
        self.root.wait_window(dialog.dialog)
        self.load_semester_goals()

    def update_semester_progress(self):
        """Обновление прогресса целей на семестр"""
        selection = self.semester_tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите цель")
            return

        item = self.semester_tree.item(selection[0])
        goal_id = item['values'][0]
        goal_type = item['values'][2]

        # Для целей типа "Количество" обновляем прогресс на основе завершённых целей
        if goal_type == "Количество":
            self.cursor.execute('SELECT COUNT(*) FROM цели WHERE статус = %s'
                                if self.db_type == 'postgres' else
                                'SELECT COUNT(*) FROM цели WHERE статус = ?',
                                ('Завершено',))
            completed = self.cursor.fetchone()[0]

            self.cursor.execute('UPDATE цель_на_семестр SET текущий_прогресс = %s WHERE id = %s'
                                if self.db_type == 'postgres' else
                                'UPDATE цель_на_семестр SET текущий_прогресс = ? WHERE id = ?',
                                (completed, goal_id))

        # Для целей типа "Поднять компетенцию" проверяем уровень компетенций
        elif goal_type == "Поднять компетенцию":
            # Здесь можно добавить логику для конкретных компетенций
            # Для примера просто увеличиваем на 1
            self.cursor.execute('SELECT текущий_прогресс FROM цель_на_семестр WHERE id = %s'
                                if self.db_type == 'postgres' else
                                'SELECT текущий_прогресс FROM цель_на_семестр WHERE id = ?',
                                (goal_id,))
            result = self.cursor.fetchone()
            current = result[0] if result else 0

            self.cursor.execute('UPDATE цель_на_семестр SET текущий_прогресс = %s WHERE id = %s'
                                if self.db_type == 'postgres' else
                                'UPDATE цель_на_семестр SET текущий_прогресс = ? WHERE id = ?',
                                (min(current + 1, item['values'][4]), goal_id))

        self.conn.commit()
        self.load_semester_goals()
        messagebox.showinfo("Успех", "Прогресс обновлен")

    def save_specialty(self):
        """Сохранение специальности"""
        self.specialty = self.specialty_var.get()
        if self.specialty:
            messagebox.showinfo("Успех", f"Специальность сохранена: {self.specialty}")

    def check_achievements(self):
        """Проверка и обновление достижений"""
        # 1. Старт - создана хотя бы одна цель
        self.cursor.execute("SELECT COUNT(*) FROM цели")
        if self.cursor.fetchone()[0] > 0:
            self.grant_achievement('start')

        # 2. Пунктуальный - три или более завершённых цели с фактической датой не позже плановой
        self.cursor.execute('''
            SELECT COUNT(*) FROM цели 
            WHERE статус = 'Завершено' 
            AND факт_дата <= план_дата
        ''')
        result = self.cursor.fetchone()
        if result and result[0] >= 3:
            self.grant_achievement('punctual')

        # 3. Многоцелевой - есть цели минимум трёх разных типов
        self.cursor.execute("SELECT COUNT(DISTINCT тип) FROM цели")
        result = self.cursor.fetchone()
        if result and result[0] >= 3:
            self.grant_achievement('multitarget')

        # 4. Навыковый рост - у одного навыка четыре или более связанных завершённых целей
        self.cursor.execute('''
            SELECT н.название, COUNT(цн.цель_id) as количество
            FROM навыки н
            JOIN цель_навыки цн ON н.id = цн.навык_id
            JOIN цели ц ON цн.цель_id = ц.id AND ц.статус = 'Завершено'
            GROUP BY н.id
            HAVING COUNT(цн.цель_id) >= 4
        ''')
        if self.cursor.fetchone():
            self.grant_achievement('skill_growth')

        # 5. Планировщик - одновременно в статусе В процессе пять или более целей
        self.cursor.execute("SELECT COUNT(*) FROM цели WHERE статус = 'В процессе'")
        result = self.cursor.fetchone()
        if result and result[0] >= 5:
            self.grant_achievement('planner')

        self.update_achievements()

    def grant_achievement(self, code):
        """Выдача достижения"""
        self.cursor.execute('UPDATE достижения SET получено = 1 WHERE код = %s AND получено = 0'
                            if self.db_type == 'postgres' else
                            'UPDATE достижения SET получено = 1 WHERE код = ? AND получено = 0',
                            (code,))
        self.conn.commit()

    def validate_date(self, date_str):
        """Проверка формата даты ГГГГ-ММ-ДД"""
        try:
            if date_str:
                datetime.strptime(date_str, '%Y-%m-%d')
            return True
        except ValueError:
            return False

    def generate_report(self):
        """Генерация отчёта в формате Word"""
        try:
            print("Начало формирования отчёта...")  # Отладочное сообщение
            doc = Document()

            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # Заголовок
            title = doc.add_heading('Индивидуальный образовательный маршрут', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.name = 'Times New Roman'
            title.runs[0].font.size = Pt(16)
            title.runs[0].bold = True

            # Разделитель
            doc.add_paragraph()

            # Раздел: Цели
            doc.add_heading('1. Учебные цели', 1)

            self.cursor.execute('''
                SELECT название, тип, статус, план_дата, факт_дата, описание 
                FROM цели ORDER BY план_дата
            ''')

            goals = self.cursor.fetchall()
            if not goals:
                doc.add_paragraph('Нет созданных целей', style='Intense Quote')
            else:
                for i, goal in enumerate(goals, 1):
                    name, type_, status, plan_date, fact_date, desc = goal

                    # Название цели как подзаголовок
                    goal_title = doc.add_heading(f'Цель {i}: {name}', 2)
                    goal_title.runs[0].font.name = 'Times New Roman'

                    # Информация о цели в таблице
                    table = doc.add_table(rows=5, cols=2)
                    table.style = 'Light Grid Accent 1'

                    # Заполняем таблицу
                    data = [
                        ('Тип цели:', type_),
                        ('Статус:', status),
                        ('Плановая дата:', plan_date if plan_date else 'Не указана'),
                        ('Фактическая дата:', fact_date if fact_date else 'Не указана'),
                        ('Дата создания:', datetime.now().strftime('%Y-%m-%d'))
                    ]

                    for row_idx, (label, value) in enumerate(data):
                        cells = table.rows[row_idx].cells
                        cells[0].text = label
                        cells[1].text = str(value)
                        # Делаем первую колонку жирной
                        cells[0].paragraphs[0].runs[0].bold = True

                    # Описание
                    # Обработка описания с поддержкой гиперссылок
                    if desc and desc.strip():
                        doc.add_heading('Описание:', 3)
                        # Разбиваем на строки и обрабатываем каждую отдельно
                        desc_lines = desc.strip().split('\n')
                        for desc_line in desc_lines:
                            self._process_line_with_hyperlinks(doc, desc_line)

                    # Разделитель между целями
                    if i < len(goals):
                        doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Раздел: Навыки
            doc.add_heading('2. Развиваемые навыки', 1)
            self.cursor.execute('''
                SELECT н.название, COUNT(DISTINCT цн.цель_id) as количество
                FROM навыки н
                LEFT JOIN цель_навыки цн ON н.id = цн.навык_id
                LEFT JOIN цели ц ON цн.цель_id = ц.id AND ц.статус = 'Завершено'
                GROUP BY н.id
                HAVING COUNT(DISTINCT цн.цель_id) > 0
                ORDER BY количество DESC, н.название
            ''')

            skills = self.cursor.fetchall()
            if skills:
                for skill, count in skills:
                    doc.add_paragraph(f'• {skill} — {count} целей', style='List Bullet')
            else:
                doc.add_paragraph('Навыки не указаны', style='Intense Quote')

            # Раздел: Компетенции
            doc.add_heading('3. Компетенции', 1)
            self.cursor.execute('''
                SELECT к.название, к.категория, 
                       COALESCE(AVG(цк.уровень), 0) as средний_уровень,
                       COUNT(цк.id) as количество_целей
                FROM компетенции к
                LEFT JOIN цель_компетенции цк ON к.id = цк.компетенция_id
                LEFT JOIN цели ц ON цк.цель_id = ц.id AND ц.статус = 'Завершено'
                GROUP BY к.id, к.название, к.категория
                ORDER BY средний_уровень DESC, количество_целей DESC
            ''')

            competencies = self.cursor.fetchall()
            if competencies:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'

                # Заголовки таблицы
                hdr_cells = table.rows[0].cells
                headers = ['Название компетенции', 'Категория', 'Средний уровень', 'Кол-во целей']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].bold = True

                # Данные
                for name, category, avg_level, count in competencies:
                    row_cells = table.add_row().cells
                    row_cells[0].text = name
                    row_cells[1].text = category if category else 'Не указана'
                    row_cells[2].text = f"{avg_level:.1f}"
                    row_cells[3].text = str(count)

                    # Подсвечиваем низкий уровень - ИСПРАВЛЕННАЯ ЧАСТЬ
                    if avg_level < 2:
                        row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Красный
                    elif avg_level < 3:
                        row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 153, 0)  # Оранжевый
            else:
                doc.add_paragraph('Компетенции не загружены', style='Intense Quote')
            # Раздел: Слабые зоны
            doc.add_heading('4. Области для развития', 1)
            weak_zones = [c for c in competencies if c[2] < 3]
            if weak_zones:
                doc.add_paragraph('Следующие компетенции требуют дополнительного внимания:')
                for name, category, avg_level, _ in weak_zones:
                    level_text = {
                        0: 'Не развивается',
                        1: 'Начальный уровень',
                        2: 'Требует улучшения'
                    }.get(int(avg_level), 'Требует улучшения')

                    doc.add_paragraph(
                        f'• {name} ({category}): {avg_level:.1f} — {level_text}',
                        style='List Bullet'
                    )
            else:
                doc.add_paragraph('Все компетенции развиты на достаточном уровне.',
                                  style='Intense Quote')

            # Раздел: Достижения
            doc.add_heading('5. Достижения', 1)
            self.cursor.execute('''
                SELECT название, описание FROM достижения WHERE получено = 1
                ORDER BY код
            ''')

            achievements = self.cursor.fetchall()
            if achievements:
                for name, desc in achievements:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{name}: ').bold = True
                    p.add_run(desc)
            else:
                doc.add_paragraph('Достижения пока не получены.', style='Intense Quote')

            # Раздел: Цели на семестр
            doc.add_heading('6. Цели на текущий семестр', 1)
            self.cursor.execute('''
                SELECT текст_цели, тип_цели, текущий_прогресс, целевой_прогресс, параметр
                FROM цель_на_семестр
                ORDER BY id
            ''')

            semester_goals = self.cursor.fetchall()
            if semester_goals:
                for goal_text, goal_type, current, target, param in semester_goals:
                    # Прогресс в процентах
                    if target > 0:
                        percentage = (current / target) * 100
                        progress_text = f" ({percentage:.0f}%)"
                    else:
                        progress_text = ""

                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{goal_text} ').bold = True

                    if goal_type == 'Поднять компетенцию' and param:
                        p.add_run(f'({param}) ')

                    p.add_run(f'— {current} из {target}{progress_text}')

                    # Добавляем прогресс-бар текстом
                    if target > 0:
                        filled = int((current / target) * 10)
                        progress_bar = '█' * filled + '░' * (10 - filled)
                        doc.add_paragraph(f'  [{progress_bar}]', style='Quote')
            else:
                doc.add_paragraph('Цели на семестр не установлены.', style='Intense Quote')

            # Раздел: Рекомендации
            doc.add_heading('7. Рекомендации', 1)

            recommendations = []

            # Рекомендации на основе компетенций
            for name, category, avg_level, count in competencies:
                if avg_level < 2:
                    recommendations.append(
                        f"• Начать развитие компетенции '{name}' через базовые курсы "
                        f"или практические задания"
                    )
                elif avg_level < 3:
                    recommendations.append(
                        f"• Углубить знания по компетенции '{name}' через продвинутые "
                        f"курсы или реальные проекты"
                    )

            # Рекомендации на основе целей
            self.cursor.execute("SELECT COUNT(*) FROM цели WHERE статус = 'В процессе'")
            in_progress = self.cursor.fetchone()[0]
            if in_progress > 5:
                recommendations.append(
                    "• Сосредоточьтесь на меньшем количестве целей для повышения эффективности"
                )
            elif in_progress < 2:
                recommendations.append(
                    "• Добавьте больше целей в работу для оптимального использования времени"
                )

            # Рекомендации на основе своевременности
            self.cursor.execute('''
                SELECT COUNT(CASE WHEN факт_дата <= план_дата THEN 1 END) * 100.0 / 
                       COUNT(CASE WHEN статус = 'Завершено' AND факт_дата IS NOT NULL THEN 1 END)
                FROM цели
                WHERE статус = 'Завершено' AND план_дата IS NOT NULL AND факт_дата IS NOT NULL
            ''')
            timely_percent = self.cursor.fetchone()[0] or 0
            if timely_percent < 80:
                recommendations.append(
                    f"• Улучшите планирование сроков (в срок выполнено {timely_percent:.0f}% целей)"
                )

            if recommendations:
                for rec in recommendations[:5]:  # Ограничиваем 5 рекомендациями
                    doc.add_paragraph(rec, style='List Bullet')
            else:
                doc.add_paragraph('Продолжайте в том же духе! Все показатели в норме.',
                                  style='Intense Quote')

            # Подпись и дата
            doc.add_page_break()
            doc.add_paragraph('\n' * 3)
            p = doc.add_paragraph()
            p.add_run('Сгенерировано: ').bold = True
            p.add_run(datetime.now().strftime('%d.%m.%Y %H:%M'))

            p = doc.add_paragraph()
            p.add_run('Студент: ').bold = True
            if self.specialty:
                p.add_run(f'{self.specialty}')

            # Сохраняем документ
            filename = f'Отчёт_ИОМ_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            doc.save(filename)
            messagebox.showinfo("Успех", f"Отчёт сохранён в файле '{filename}'")

            # Автоматическое открытие файла (опционально)
            try:
                os.startfile(filename)
            except:
                pass

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сформировать отчёт: {str(e)}")

    def format_description(self, doc, text):
        """Форматирование описания для Word"""
        lines = text.split('\n')
        current_paragraph = None

        for line in lines:
            line = line.rstrip()
            if not line:
                continue

            if line.startswith('# '):
                # Заголовок
                doc.add_heading(line[2:], 2)
                current_paragraph = None
            elif line.startswith('## '):
                # Подзаголовок
                doc.add_heading(line[3:], 3)
                current_paragraph = None
            elif line.startswith('- ') or line.startswith('* '):
                # Маркированный список
                if current_paragraph and hasattr(current_paragraph,
                                                 'style') and current_paragraph.style.name == 'List Bullet':
                    # Продолжаем текущий список
                    current_paragraph.add_run(f'\n{line[2:]}')
                else:
                    # Начинаем новый список
                    current_paragraph = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('1) '):
                # Нумерованный список
                if current_paragraph and hasattr(current_paragraph,
                                                 'style') and current_paragraph.style.name == 'List Number':
                    current_paragraph.add_run(f'\n{line[3:]}')
                else:
                    current_paragraph = doc.add_paragraph(line[3:], style='List Number')
            elif '**' in line:
                # Обработка жирного текста
                parts = line.split('**')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # Нечётные части - жирный текст
                        run.bold = True
                current_paragraph = p
            elif '*' in line or '_' in line:
                # Обработка курсива
                p = doc.add_paragraph()
                parts_star = line.split('*')
                parts_underscore = line.split('_')

                # Используем разделение по звёздочкам
                if len(parts_star) > 1:
                    for i, part in enumerate(parts_star):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.italic = True
                # Или по подчёркиваниям
                elif len(parts_underscore) > 1:
                    for i, part in enumerate(parts_underscore):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.italic = True
                else:
                    p.add_run(line)
                current_paragraph = p
            else:
                # Обработка строк, которые могут содержать гиперссылки
                self._process_line_with_hyperlinks(doc, line)
                current_paragraph = None

    def _process_line_with_hyperlinks(self, doc, line):
        """Обработка строки с гиперссылками для Word документа"""
        import re

        # Шаблон для поиска гиперссылок [текст](url)
        link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'

        # Находим все гиперссылки в строке
        matches = list(re.finditer(link_pattern, line))

        if not matches:
            # Если нет гиперссылок, добавляем как обычный текст
            if line.strip():
                doc.add_paragraph(line)
            return

        # Создаем параграф для всей строки
        p = doc.add_paragraph()
        last_pos = 0

        for match in matches:
            # Добавляем текст перед гиперссылкой
            if match.start() > last_pos:
                p.add_run(line[last_pos:match.start()])

            # Извлекаем текст и URL гиперссылки
            link_text = match.group(1)
            url = match.group(2)

            try:
                # Добавляем гиперссылку в Word
                from docx.oxml.shared import qn

                # Создаем элемент гиперссылки
                hyperlink = doc._element._part._element.body.add_hyperlink(url, link_text)

                # Создаем run для текста гиперссылки
                run = doc._element._part._element.body.add_run(link_text)

                # Применяем стиль к гиперссылке (синий цвет, подчеркивание)
                run.rPr = doc._element._part._element.body.add_rPr()
                run.rPr.color = doc._element._part._element.body.add_color()
                run.rPr.color.val = "0000FF"  # Синий цвет
                run.rPr.u = doc._element._part._element.body.add_u()
                run.rPr.u.val = "single"  # Подчеркивание

            except Exception as e:
                # Если не удалось создать гиперссылку, добавляем как обычный текст
                print(f"Ошибка создания гиперссылки: {e}")
                run = p.add_run(f'{link_text} ({url})')
                run.font.color.rgb = RGBColor(0, 0, 255)  # Синий цвет
                run.underline = True

            last_pos = match.end()

        # Добавляем оставшийся текст после последней гиперссылки
        if last_pos < len(line):
            p.add_run(line[last_pos:])

    def preview_markdown(self, text, preview_widget):
        """Предпросмотр разметки"""
        preview_widget.delete(1.0, tk.END)
        lines = text.split('\n')

        for line in lines:
            line = line.rstrip()
            if line.startswith('# '):
                preview_widget.insert(tk.END, f"ЗАГОЛОВОК: {line[2:]}\n")
            elif line.startswith('- ') or line.startswith('* '):
                preview_widget.insert(tk.END, f"• {line[2:]}\n")
            elif '**' in line:
                # Заменяем **текст** на заглавные буквы
                parts = line.split('**')
                formatted = ''
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        formatted += part.upper()
                    else:
                        formatted += part
                preview_widget.insert(tk.END, f"{formatted}\n")
            else:
                # Обработка гиперссылок
                import re
                link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'

                # Заменяем все найденные ссылки
                result_line = line
                matches = list(re.finditer(link_pattern, line))

                if matches:
                    # Заменяем с конца, чтобы не сбивать индексы
                    for match in reversed(matches):
                        link_text = match.group(1)
                        url = match.group(2)
                        result_line = (result_line[:match.start()] +
                                       f"[ссылка: {link_text} -> {url}]" +
                                       result_line[match.end():])

                preview_widget.insert(tk.END, f"{result_line}\n")

class GoalDialog:
    """Диалог для добавления/редактирования цели"""

    def __init__(self, parent, title, app, goal_id=None):
        self.app = app
        self.goal_id = goal_id
        self.skills_vars = []

        self.dialog = tk.Toplevel(parent)
        self.dialog.title(title)
        self.dialog.geometry("900x700")
        self.dialog.transient(parent)
        self.dialog.grab_set()

        self.create_widgets()
        if goal_id:
            self.load_goal_data()

    def create_widgets(self):
        """Создание виджетов диалога"""
        main_frame = ttk.Frame(self.dialog)
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)

        # Левая часть - форма
        left_frame = ttk.Frame(main_frame)
        left_frame.pack(side='left', fill='both', expand=True, padx=(0, 5))

        # Название
        ttk.Label(left_frame, text="Название:").grid(row=0, column=0, sticky='w', pady=5)
        self.name_entry = ttk.Entry(left_frame, width=40)
        self.name_entry.grid(row=0, column=1, sticky='ew', pady=5, padx=(5, 0))

        # Тип
        ttk.Label(left_frame, text="Тип:").grid(row=1, column=0, sticky='w', pady=5)
        self.type_var = tk.StringVar()
        types = ["Курс", "Проект", "Самостоятельная работа", "Конференция", "Сертификация", "Другое"]
        self.type_combo = ttk.Combobox(left_frame, textvariable=self.type_var, values=types, state='readonly')
        self.type_combo.grid(row=1, column=1, sticky='ew', pady=5, padx=(5, 0))

        # Статус
        ttk.Label(left_frame, text="Статус:").grid(row=2, column=0, sticky='w', pady=5)
        self.status_var = tk.StringVar()
        statuses = ["Планируется", "В процессе", "Завершено", "Отменено"]
        self.status_combo = ttk.Combobox(left_frame, textvariable=self.status_var, values=statuses, state='readonly')
        self.status_combo.grid(row=2, column=1, sticky='ew', pady=5, padx=(5, 0))

        # Даты
        ttk.Label(left_frame, text="Плановая дата (ГГГГ-ММ-ДД):").grid(row=3, column=0, sticky='w', pady=5)
        self.plan_date_entry = ttk.Entry(left_frame, width=15)
        self.plan_date_entry.grid(row=3, column=1, sticky='w', pady=5, padx=(5, 0))

        ttk.Label(left_frame, text="Фактическая дата (ГГГГ-ММ-ДД):").grid(row=4, column=0, sticky='w', pady=5)
        self.fact_date_entry = ttk.Entry(left_frame, width=15)
        self.fact_date_entry.grid(row=4, column=1, sticky='w', pady=5, padx=(5, 0))

        # Навыки
        ttk.Label(left_frame, text="Навыки (до 3-х):").grid(row=5, column=0, sticky='w', pady=5)
        skills_frame = ttk.Frame(left_frame)
        skills_frame.grid(row=5, column=1, sticky='ew', pady=5, padx=(5, 0))

        for i in range(3):
            var = tk.StringVar()
            self.skills_vars.append(var)
            combo = ttk.Combobox(skills_frame, textvariable=var,
                                 values=self.app.skills_autocomplete)
            combo.pack(fill='x', pady=2)

        # Компетенции
        ttk.Label(left_frame, text="Компетенции:").grid(row=6, column=0, sticky='w', pady=5)
        comp_frame = ttk.Frame(left_frame)
        comp_frame.grid(row=6, column=1, sticky='ew', pady=5, padx=(5, 0))

        self.app.cursor.execute("SELECT id, название FROM компетенции")
        self.competencies = self.app.cursor.fetchall()
        self.comp_vars = []
        self.level_vars = []

        for i, (comp_id, comp_name) in enumerate(self.competencies[:3]):  # Ограничиваем 3 компетенциями
            var = tk.BooleanVar()
            self.comp_vars.append((var, comp_id))

            check = ttk.Checkbutton(comp_frame, text=comp_name, variable=var)
            check.grid(row=i, column=0, sticky='w', pady=2)

            level_var = tk.StringVar(value="0")
            level_combo = ttk.Combobox(comp_frame, textvariable=level_var,
                                       values=['0', '1', '2', '3', '4', '5'],
                                       width=5, state='readonly')
            level_combo.grid(row=i, column=1, sticky='w', padx=(10, 0), pady=2)
            self.level_vars.append(level_var)

        # Описание
        ttk.Label(left_frame, text="Описание:").grid(row=7, column=0, sticky='nw', pady=5)
        self.desc_text = tk.Text(left_frame, width=40, height=10)
        self.desc_text.grid(row=7, column=1, sticky='nsew', pady=5, padx=(5, 0))

        # Правая часть - предпросмотр
        right_frame = ttk.Frame(main_frame)
        right_frame.pack(side='right', fill='both', expand=True, padx=(5, 0))

        ttk.Label(right_frame, text="Предпросмотр:").pack(anchor='w', pady=(0, 5))
        self.preview_text = tk.Text(right_frame, width=40, height=30, state='normal')
        self.preview_text.pack(fill='both', expand=True)

        # Кнопки
        button_frame = ttk.Frame(self.dialog)
        button_frame.pack(fill='x', padx=10, pady=10)

        ttk.Button(button_frame, text="Предпросмотр",
                   command=self.update_preview).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Сохранить",
                   command=self.save_goal).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Отмена",
                   command=self.dialog.destroy).pack(side='right', padx=5)

        # Привязка события для автоматического предпросмотра
        self.desc_text.bind('<KeyRelease>', lambda e: self.update_preview())

    def load_goal_data(self):
        """Загрузка данных цели для редактирования"""
        self.app.cursor.execute('SELECT название, тип, статус, план_дата, факт_дата, описание FROM цели WHERE id = %s'
                                if self.app.db_type == 'postgres' else
                                'SELECT название, тип, статус, план_дата, факт_дата, описание FROM цели WHERE id = ?',
                                (self.goal_id,))

        goal = self.app.cursor.fetchone()
        if goal:
            self.name_entry.insert(0, goal[0])
            self.type_var.set(goal[1])
            self.status_var.set(goal[2])
            if goal[3]:
                self.plan_date_entry.insert(0, goal[3])
            if goal[4]:
                self.fact_date_entry.insert(0, goal[4])
            if goal[5]:
                self.desc_text.insert(1.0, goal[5])
                self.update_preview()

        # Загружаем навыки
        self.app.cursor.execute('''
            SELECT н.название 
            FROM навыки н
            JOIN цель_навыки цн ON н.id = цн.навык_id
            WHERE цн.цель_id = %s
        ''' if self.app.db_type == 'postgres' else '''
            SELECT н.название 
            FROM навыки н
            JOIN цель_навыки цн ON н.id = цн.навык_id
            WHERE цн.цель_id = ?
        ''', (self.goal_id,))

        skills = self.app.cursor.fetchall()
        for i, skill in enumerate(skills[:3]):
            if i < len(self.skills_vars):
                self.skills_vars[i].set(skill[0])

        # Загружаем компетенции
        self.app.cursor.execute('SELECT компетенция_id, уровень FROM цель_компетенции WHERE цель_id = %s'
                                if self.app.db_type == 'postgres' else
                                'SELECT компетенция_id, уровень FROM цель_компетенции WHERE цель_id = ?',
                                (self.goal_id,))

        comps = self.app.cursor.fetchall()
        for comp_id, level in comps:
            for i, (_, c_id) in enumerate(self.comp_vars):
                if c_id == comp_id and i < len(self.level_vars):
                    self.comp_vars[i][0].set(True)
                    self.level_vars[i].set(str(level))

    def update_preview(self):
        """Обновление предпросмотра"""
        text = self.desc_text.get(1.0, tk.END)
        self.app.preview_markdown(text, self.preview_text)

    def save_goal(self):
        """Сохранение цели"""
        # Проверяем обязательные поля
        if not self.name_entry.get().strip():
            messagebox.showwarning("Предупреждение", "Введите название цели")
            return

        if not self.type_var.get():
            messagebox.showwarning("Предупреждение", "Выберите тип цели")
            return

        if not self.status_var.get():
            messagebox.showwarning("Предупреждение", "Выберите статус цели")
            return

        # Проверяем даты
        if self.plan_date_entry.get().strip() and not self.app.validate_date(self.plan_date_entry.get().strip()):
            messagebox.showwarning("Ошибка", "Неверный формат плановой даты. Используйте ГГГГ-ММ-ДД")
            return

        if self.fact_date_entry.get().strip() and not self.app.validate_date(self.fact_date_entry.get().strip()):
            messagebox.showwarning("Ошибка", "Неверный формат фактической даты. Используйте ГГГГ-ММ-ДД")
            return

        # Сохраняем цель
        if self.goal_id:
            # Обновляем существующую цель
            self.app.cursor.execute('''
                UPDATE цели SET 
                    название = %s, тип = %s, статус = %s, 
                    план_дата = %s, факт_дата = %s, описание = %s
                WHERE id = %s
            ''' if self.app.db_type == 'postgres' else '''
                UPDATE цели SET 
                    название = ?, тип = ?, статус = ?, 
                    план_дата = ?, факт_дата = ?, описание = ?
                WHERE id = ?
            ''', (
                self.name_entry.get().strip(),
                self.type_var.get(),
                self.status_var.get(),
                self.plan_date_entry.get().strip() or None,
                self.fact_date_entry.get().strip() or None,
                self.desc_text.get(1.0, tk.END).strip(),
                self.goal_id
            ))
        else:
            # Добавляем новую цель
            self.app.cursor.execute('''
                INSERT INTO цели (название, тип, статус, план_дата, факт_дата, описание)
                VALUES (%s, %s, %s, %s, %s, %s)
            ''' if self.app.db_type == 'postgres' else '''
                INSERT INTO цели (название, тип, статус, план_дата, факт_дата, описание)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (
                self.name_entry.get().strip(),
                self.type_var.get(),
                self.status_var.get(),
                self.plan_date_entry.get().strip() or None,
                self.fact_date_entry.get().strip() or None,
                self.desc_text.get(1.0, tk.END).strip()
            ))
            if self.app.db_type == 'postgres':
                self.goal_id = self.app.cursor.fetchone()[0] if self.app.cursor.description else None
            else:
                self.goal_id = self.app.cursor.lastrowid

        # Удаляем старые связи с навыками
        self.app.cursor.execute("DELETE FROM цель_навыки WHERE цель_id = %s"
                                if self.app.db_type == 'postgres' else
                                "DELETE FROM цель_навыки WHERE цель_id = ?",
                                (self.goal_id,))

        # Добавляем навыки
        for var in self.skills_vars:
            skill_name = var.get().strip()
            if skill_name:
                # Проверяем, есть ли навык в базе
                self.app.cursor.execute("SELECT id FROM навыки WHERE название = %s"
                                        if self.app.db_type == 'postgres' else
                                        "SELECT id FROM навыки WHERE название = ?",
                                        (skill_name,))
                skill_row = self.app.cursor.fetchone()

                if not skill_row:
                    # Добавляем новый навык
                    self.app.cursor.execute("INSERT INTO навыки (название) VALUES (%s)"
                                            if self.app.db_type == 'postgres' else
                                            "INSERT INTO навыки (название) VALUES (?)",
                                            (skill_name,))
                    if self.app.db_type == 'postgres':
                        self.app.cursor.execute("SELECT LASTVAL()")
                        skill_id = self.app.cursor.fetchone()[0]
                    else:
                        skill_id = self.app.cursor.lastrowid
                else:
                    skill_id = skill_row[0]

                # Связываем цель с навыком
                self.app.cursor.execute(
                    "INSERT INTO цель_навыки (цель_id, навык_id) VALUES (%s, %s)"
                    if self.app.db_type == 'postgres' else
                    "INSERT INTO цель_навыки (цель_id, навык_id) VALUES (?, ?)",
                    (self.goal_id, skill_id)
                )

        # Удаляем старые связи с компетенциями
        self.app.cursor.execute("DELETE FROM цель_компетенции WHERE цель_id = %s"
                                if self.app.db_type == 'postgres' else
                                "DELETE FROM цель_компетенции WHERE цель_id = ?",
                                (self.goal_id,))

        # Добавляем компетенции
        for (var, comp_id), level_var in zip(self.comp_vars, self.level_vars):
            if var.get():  # Чекбокс отмечен
                level = int(level_var.get())
                self.app.cursor.execute('''
                    INSERT INTO цель_компетенции (цель_id, компетенция_id, уровень)
                    VALUES (%s, %s, %s)
                ''' if self.app.db_type == 'postgres' else '''
                    INSERT INTO цель_компетенции (цель_id, компетенция_id, уровень)
                    VALUES (?, ?, ?)
                ''', (self.goal_id, comp_id, level))

        self.app.conn.commit()
        self.app.load_skills_autocomplete()
        self.dialog.destroy()


class SemesterGoalDialog:
    """Диалог для добавления цели на семестр"""

    def __init__(self, parent, app):
        self.app = app
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Добавить цель на семестр")
        self.dialog.geometry("500x300")
        self.dialog.transient(parent)
        self.dialog.grab_set()

        self.create_widgets()

    def create_widgets(self):
        """Создание виджетов диалога"""
        frame = ttk.Frame(self.dialog)
        frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Текст цели
        ttk.Label(frame, text="Текст цели:").grid(row=0, column=0, sticky='w', pady=5)
        self.goal_entry = ttk.Entry(frame, width=40)
        self.goal_entry.grid(row=0, column=1, sticky='ew', pady=5, padx=(5, 0))

        # Тип цели
        ttk.Label(frame, text="Тип цели:").grid(row=1, column=0, sticky='w', pady=5)
        self.type_var = tk.StringVar()
        types = ["Количество", "Поднять компетенцию"]
        self.type_combo = ttk.Combobox(frame, textvariable=self.type_var, values=types, state='readonly')
        self.type_combo.grid(row=1, column=1, sticky='ew', pady=5, padx=(5, 0))
        self.type_combo.set("Количество")

        # Параметр (для целей типа "Поднять компетенцию")
        ttk.Label(frame, text="Параметр:").grid(row=2, column=0, sticky='w', pady=5)
        self.param_entry = ttk.Entry(frame, width=40)
        self.param_entry.grid(row=2, column=1, sticky='ew', pady=5, padx=(5, 0))

        # Целевой прогресс
        ttk.Label(frame, text="Целевой прогресс:").grid(row=3, column=0, sticky='w', pady=5)
        self.target_spinbox = tk.Spinbox(frame, from_=1, to=100, width=10)
        self.target_spinbox.grid(row=3, column=1, sticky='w', pady=5, padx=(5, 0))
        self.target_spinbox.delete(0, tk.END)
        self.target_spinbox.insert(0, "3")

        # Кнопки
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=20)

        ttk.Button(button_frame, text="Сохранить",
                   command=self.save_goal).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Отмена",
                   command=self.dialog.destroy).pack(side='right', padx=5)

    def save_goal(self):
        """Сохранение цели на семестр"""
        goal_text = self.goal_entry.get().strip()
        if not goal_text:
            messagebox.showwarning("Предупреждение", "Введите текст цели")
            return

        try:
            target = int(self.target_spinbox.get())
        except ValueError:
            messagebox.showwarning("Предупреждение", "Введите число для целевого прогресса")
            return

        self.app.cursor.execute('''
            INSERT INTO цель_на_семестр (текст_цели, тип_цели, параметр, целевой_прогресс)
            VALUES (%s, %s, %s, %s)
        ''' if self.app.db_type == 'postgres' else '''
            INSERT INTO цель_на_семестр (текст_цели, тип_цели, параметр, целевой_прогресс)
            VALUES (?, ?, ?, ?)
        ''', (
            goal_text,
            self.type_var.get(),
            self.param_entry.get().strip() or None,
            target
        ))

        self.app.conn.commit()
        self.dialog.destroy()


def main():
    """Основная функция"""
    root = tk.Tk()
    app = EducationalPlanner(root)
    root.mainloop()


if __name__ == "__main__":
    main()