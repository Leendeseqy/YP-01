import tkinter as tk
from tkinter import ttk, messagebox
import psycopg2
from psycopg2 import sql
from datetime import datetime
from docx import Document
import json
import os

DB_LOGIN = {
    'host': 'localhost',
    'database': 'postgres',
    'user': 'postgres',
    'password': '1111',
    'port': '5432'
}


class PortfolioApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Система портфолио и компетенций")
        self.root.geometry("1200x700")

        self.current_user_id = 1

        try:
            self.initialize_database()
            self.load_competencies()
        except Exception as e:
            messagebox.showerror("Ошибка базы данных", f"Не удалось подключиться к базе данных:\n{str(e)}")
            self.root.destroy()
            return

        self.create_menu()
        self.create_main_interface()

        self.update_statistics()

    def get_connection(self):
        return psycopg2.connect(**DB_LOGIN)

    def initialize_database(self):
        self.conn = self.get_connection()
        self.cursor = self.conn.cursor()

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS entries (
                id SERIAL PRIMARY KEY,
                title TEXT NOT NULL,
                type TEXT NOT NULL,
                date DATE NOT NULL,
                description TEXT,
                coauthors TEXT,
                user_id INTEGER
            )
        ''')

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS keywords (
                id SERIAL PRIMARY KEY,
                keyword TEXT UNIQUE NOT NULL
            )
        ''')

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS entry_keywords (
                entry_id INTEGER REFERENCES entries(id) ON DELETE CASCADE,
                keyword_id INTEGER REFERENCES keywords(id) ON DELETE CASCADE,
                PRIMARY KEY (entry_id, keyword_id)
            )
        ''')

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS achievements (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                description TEXT,
                user_id INTEGER,
                unlocked_date DATE
            )
        ''')

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS competencies (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                category TEXT
            )
        ''')

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS entry_competencies (
                entry_id INTEGER REFERENCES entries(id) ON DELETE CASCADE,
                competency_id INTEGER REFERENCES competencies(id) ON DELETE CASCADE,
                level INTEGER CHECK (level >= 1 AND level <= 5),
                PRIMARY KEY (entry_id, competency_id)
            )
        ''')

        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS goals (
                id SERIAL PRIMARY KEY,
                description TEXT NOT NULL,
                target_value INTEGER,
                current_value INTEGER,
                deadline DATE,
                user_id INTEGER
            )
        ''')

        self.conn.commit()

        self.insert_default_competencies()

    def insert_default_competencies(self):
        self.cursor.execute("SELECT COUNT(*) FROM competencies")
        if self.cursor.fetchone()[0] == 0:
            default_competencies = [
                ("Программирование", "Технические"),
                ("Работа с БД", "Технические"),
                ("Анализ данных", "Технические"),
                ("Проектная деятельность", "Профессиональные"),
                ("Научная работа", "Профессиональные"),
                ("Публикации", "Профессиональные"),
                ("Презентация результатов", "Коммуникативные"),
                ("Командная работа", "Коммуникативные"),
                ("Самоорганизация", "Личные")
            ]
            insert_query = sql.SQL("INSERT INTO competencies (name, category) VALUES {}").format(
                sql.SQL(', ').join([
                    sql.SQL("({}, {})").format(sql.Literal(name), sql.Literal(category))
                    for name, category in default_competencies
                ])
            )
            self.cursor.execute(insert_query)
            self.conn.commit()

    def load_competencies(self):
        self.competencies = {}
        self.cursor.execute("SELECT id, name, category FROM competencies")
        for row in self.cursor.fetchall():
            self.competencies[row[0]] = {"name": row[1], "category": row[2]}

    def create_menu(self):
        menubar = tk.Menu(self.root)

        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Экспорт в Word", command=self.export_to_word)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self.root.quit)
        menubar.add_cascade(label="Файл", menu=file_menu)

        self.root.config(menu=menubar)

    def create_main_interface(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.create_add_entry_tab()
        self.create_view_tab()
        self.create_research_map_tab()
        self.create_achievements_tab()
        self.create_competencies_tab()
        self.create_goals_tab()

    def create_add_entry_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Добавить запись")

        ttk.Label(frame, text="Название:").grid(row=0, column=0, sticky=tk.W, padx=10, pady=5)
        self.title_entry = ttk.Entry(frame, width=50)
        self.title_entry.grid(row=0, column=1, padx=10, pady=5)

        ttk.Label(frame, text="Тип:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)
        self.type_combo = ttk.Combobox(frame, values=["Проект", "Публикация", "Конференция", "Практика", "Грант"],
                                       state="readonly")
        self.type_combo.grid(row=1, column=1, padx=10, pady=5)

        ttk.Label(frame, text="Дата (ГГГГ-ММ-ДД):").grid(row=2, column=0, sticky=tk.W, padx=10, pady=5)
        self.date_entry = ttk.Entry(frame, width=20)
        self.date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
        self.date_entry.grid(row=2, column=1, padx=10, pady=5)

        ttk.Label(frame, text="Описание:").grid(row=3, column=0, sticky=tk.NW, padx=10, pady=5)
        self.description_text = tk.Text(frame, width=50, height=10)
        self.description_text.grid(row=3, column=1, padx=10, pady=5)

        ttk.Label(frame, text="Соавторы (через запятую):").grid(row=4, column=0, sticky=tk.W, padx=10, pady=5)
        self.coauthors_entry = ttk.Entry(frame, width=50)
        self.coauthors_entry.grid(row=4, column=1, padx=10, pady=5)

        ttk.Label(frame, text="Ключевые слова (до 5):").grid(row=5, column=0, sticky=tk.W, padx=10, pady=5)
        self.keywords_frame = ttk.Frame(frame)
        self.keywords_frame.grid(row=5, column=1, padx=10, pady=5)

        self.keyword_entries = []
        for i in range(5):
            entry = ttk.Entry(self.keywords_frame, width=20)
            entry.grid(row=0, column=i, padx=2)
            self.keyword_entries.append(entry)

        ttk.Label(frame, text="Компетенции (до 3):").grid(row=6, column=0, sticky=tk.W, padx=10, pady=5)
        self.competencies_frame = ttk.Frame(frame)
        self.competencies_frame.grid(row=6, column=1, padx=10, pady=5)

        self.competency_vars = []
        self.level_combos = []

        self.cursor.execute("SELECT id, name FROM competencies")
        competencies_list = [f"{row[0]}: {row[1]}" for row in self.cursor.fetchall()]

        for i in range(3):
            var = tk.StringVar()
            combo = ttk.Combobox(self.competencies_frame, textvariable=var,
                                 values=competencies_list, width=25, state="readonly")
            combo.grid(row=i, column=0, padx=5, pady=2)

            level_combo = ttk.Combobox(self.competencies_frame, values=["1", "2", "3", "4", "5"],
                                       width=5, state="readonly")
            level_combo.grid(row=i, column=1, padx=5, pady=2)

            self.competency_vars.append(var)
            self.level_combos.append(level_combo)

        ttk.Button(frame, text="Добавить запись", command=self.add_entry).grid(row=7, column=1, pady=20)

    def create_view_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Мои записи")

        columns = ("ID", "Название", "Тип", "Дата")
        self.tree = ttk.Treeview(frame, columns=columns, show="headings", height=20)

        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=150)

        scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.grid(row=0, column=0, padx=10, pady=10, sticky=tk.NSEW)
        scrollbar.grid(row=0, column=1, sticky=tk.NS)

        frame.grid_rowconfigure(0, weight=1)
        frame.grid_columnconfigure(0, weight=1)

        self.load_entries()

    def create_research_map_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Исследовательская карта")

        ttk.Label(frame, text="Ключевые слова:", font=("Arial", 12, "bold")).grid(row=0, column=0, sticky=tk.W, padx=10,
                                                                                  pady=(10, 5))
        self.keywords_text = tk.Text(frame, width=40, height=15, state=tk.DISABLED)
        self.keywords_text.grid(row=1, column=0, padx=10, pady=5, sticky=tk.W)

        ttk.Label(frame, text="Соавторы:", font=("Arial", 12, "bold")).grid(row=0, column=1, sticky=tk.W, padx=10,
                                                                            pady=(10, 5))
        self.coauthors_text = tk.Text(frame, width=40, height=15, state=tk.DISABLED)
        self.coauthors_text.grid(row=1, column=1, padx=10, pady=5, sticky=tk.W)

    def create_achievements_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Достижения")

        self.achievements_text = tk.Text(frame, width=80, height=25, state=tk.DISABLED)
        scrollbar = ttk.Scrollbar(frame, command=self.achievements_text.yview)
        self.achievements_text.configure(yscrollcommand=scrollbar.set)

        self.achievements_text.grid(row=0, column=0, padx=10, pady=10, sticky=tk.NSEW)
        scrollbar.grid(row=0, column=1, sticky=tk.NS)

        frame.grid_rowconfigure(0, weight=1)
        frame.grid_columnconfigure(0, weight=1)

    def create_competencies_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Мои компетенции")

        left_frame = ttk.Frame(frame)
        left_frame.grid(row=0, column=0, padx=10, pady=10, sticky=tk.NSEW)

        ttk.Label(left_frame, text="Средний уровень компетенций:", font=("Arial", 11, "bold")).grid(row=0, column=0,
                                                                                                    sticky=tk.W, pady=5)
        self.competencies_text = tk.Text(left_frame, width=40, height=12, state=tk.DISABLED)
        self.competencies_text.grid(row=1, column=0, pady=5)

        ttk.Label(left_frame, text="Слабые зоны (уровень < 3):", font=("Arial", 11, "bold")).grid(row=2, column=0,
                                                                                                  sticky=tk.W, pady=5)
        self.weak_zones_text = tk.Text(left_frame, width=40, height=6, state=tk.DISABLED)
        self.weak_zones_text.grid(row=3, column=0, pady=5)

        right_frame = ttk.Frame(frame)
        right_frame.grid(row=0, column=1, padx=10, pady=10, sticky=tk.NSEW)

        ttk.Label(right_frame, text="Рекомендации:", font=("Arial", 11, "bold")).grid(row=0, column=0, sticky=tk.W,
                                                                                      pady=5)
        self.recommendations_text = tk.Text(right_frame, width=40, height=20, state=tk.DISABLED)
        self.recommendations_text.grid(row=1, column=0, pady=5)

        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(1, weight=1)

    def create_goals_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Цели на семестр")

        ttk.Label(frame, text="Новая цель:").grid(row=0, column=0, sticky=tk.W, padx=10, pady=10)
        self.goal_entry = ttk.Entry(frame, width=50)
        self.goal_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(frame, text="Целевое значение:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)
        self.target_entry = ttk.Entry(frame, width=10)
        self.target_entry.grid(row=1, column=1, sticky=tk.W, padx=10, pady=5)

        ttk.Button(frame, text="Добавить цель", command=self.add_goal).grid(row=2, column=1, sticky=tk.W, padx=10,
                                                                            pady=10)

        ttk.Label(frame, text="Текущие цели:", font=("Arial", 11, "bold")).grid(row=3, column=0, columnspan=2,
                                                                                sticky=tk.W, padx=10, pady=(20, 5))

        self.goals_text = tk.Text(frame, width=80, height=15, state=tk.DISABLED)
        scrollbar = ttk.Scrollbar(frame, command=self.goals_text.yview)
        self.goals_text.configure(yscrollcommand=scrollbar.set)

        self.goals_text.grid(row=4, column=0, columnspan=2, padx=10, pady=5, sticky=tk.NSEW)
        scrollbar.grid(row=4, column=2, sticky=tk.NS)

        frame.grid_rowconfigure(4, weight=1)
        frame.grid_columnconfigure(1, weight=1)

        self.load_goals()

    def add_entry(self):
        title = self.title_entry.get().strip()
        entry_type = self.type_combo.get()
        date = self.date_entry.get().strip()
        description = self.description_text.get("1.0", tk.END).strip()
        coauthors = self.coauthors_entry.get().strip()

        if not title or not entry_type or not date:
            messagebox.showerror("Ошибка", "Заполните обязательные поля: Название, Тип, Дата")
            return

        try:
            datetime.strptime(date, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("Ошибка", "Неверный формат даты. Используйте ГГГГ-ММ-ДД")
            return

        keywords = [kw.get().strip() for kw in self.keyword_entries if kw.get().strip()]

        competencies = []
        for i in range(3):
            comp_val = self.competency_vars[i].get()
            level_val = self.level_combos[i].get()
            if comp_val and level_val:
                comp_id = int(comp_val.split(":")[0])
                level = int(level_val)
                competencies.append((comp_id, level))

        try:
            self.cursor.execute(
                "INSERT INTO entries (title, type, date, description, coauthors, user_id) VALUES (%s, %s, %s, %s, %s, %s) RETURNING id",
                (title, entry_type, date, description, coauthors, self.current_user_id)
            )
            entry_id = self.cursor.fetchone()[0]

            for keyword in keywords:
                self.cursor.execute("INSERT INTO keywords (keyword) VALUES (%s) ON CONFLICT (keyword) DO NOTHING",
                                    (keyword,))
                self.cursor.execute("SELECT id FROM keywords WHERE keyword = %s", (keyword,))
                keyword_row = self.cursor.fetchone()
                if keyword_row:
                    keyword_id = keyword_row[0]
                    self.cursor.execute(
                        "INSERT INTO entry_keywords (entry_id, keyword_id) VALUES (%s, %s)",
                        (entry_id, keyword_id)
                    )

            for comp_id, level in competencies:
                self.cursor.execute(
                    "INSERT INTO entry_competencies (entry_id, competency_id, level) VALUES (%s, %s, %s)",
                    (entry_id, comp_id, level)
                )

            self.conn.commit()

            self.clear_entry_form()
            self.load_entries()
            self.update_statistics()
            self.check_achievements()
            messagebox.showinfo("Успех", "Запись успешно добавлена")

        except Exception as e:
            self.conn.rollback()
            messagebox.showerror("Ошибка базы данных", f"Не удалось добавить запись:\n{str(e)}")

    def clear_entry_form(self):
        self.title_entry.delete(0, tk.END)
        self.type_combo.set('')
        self.date_entry.delete(0, tk.END)
        self.date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
        self.description_text.delete("1.0", tk.END)
        self.coauthors_entry.delete(0, tk.END)
        for entry in self.keyword_entries:
            entry.delete(0, tk.END)
        for var in self.competency_vars:
            var.set('')
        for combo in self.level_combos:
            combo.set('')

    def load_entries(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        self.cursor.execute("SELECT id, title, type, date FROM entries WHERE user_id = %s ORDER BY date DESC",
                            (self.current_user_id,))

        for row in self.cursor.fetchall():
            self.tree.insert("", tk.END, values=(row[0], row[1], row[2], row[3]))

    def update_statistics(self):
        self.cursor.execute("""
            SELECT k.keyword, COUNT(ek.entry_id) as count
            FROM keywords k
            JOIN entry_keywords ek ON k.id = ek.keyword_id
            JOIN entries e ON ek.entry_id = e.id
            WHERE e.user_id = %s
            GROUP BY k.keyword
            ORDER BY count DESC
        """, (self.current_user_id,))

        keywords_text_content = ""
        for row in self.cursor.fetchall():
            keywords_text_content += f"{row[0]} — {row[1]} записей\n"

        self.keywords_text.config(state=tk.NORMAL)
        self.keywords_text.delete("1.0", tk.END)
        self.keywords_text.insert("1.0", keywords_text_content if keywords_text_content else "Нет данных")
        self.keywords_text.config(state=tk.DISABLED)

        self.cursor.execute("""
            SELECT e.coauthors
            FROM entries e
            WHERE e.user_id = %s AND e.coauthors IS NOT NULL AND e.coauthors != ''
        """, (self.current_user_id,))

        coauthors_dict = {}
        for row in self.cursor.fetchall():
            coauthors = [c.strip() for c in row[0].split(",") if c.strip()]
            for coauthor in coauthors:
                coauthors_dict[coauthor] = coauthors_dict.get(coauthor, 0) + 1

        coauthors_text_content = ""
        for coauthor, count in sorted(coauthors_dict.items(), key=lambda x: x[1], reverse=True):
            coauthors_text_content += f"{coauthor} — {count} работ\n"

        self.coauthors_text.config(state=tk.NORMAL)
        self.coauthors_text.delete("1.0", tk.END)
        self.coauthors_text.insert("1.0", coauthors_text_content if coauthors_text_content else "Нет данных")
        self.coauthors_text.config(state=tk.DISABLED)

        self.update_competencies_dashboard()
        self.update_achievements()
        self.load_goals()

    def update_competencies_dashboard(self):
        self.cursor.execute("""
            SELECT c.name, CAST(AVG(ec.level) AS DECIMAL(10,2)) as avg_level
            FROM competencies c
            LEFT JOIN entry_competencies ec ON c.id = ec.competency_id
            LEFT JOIN entries e ON ec.entry_id = e.id AND e.user_id = %s
            GROUP BY c.id, c.name
            HAVING AVG(ec.level) IS NOT NULL
            ORDER BY c.category, c.name
        """, (self.current_user_id,))

        competencies_text_content = ""
        weak_zones_content = ""
        recommendations_content = ""

        for row in self.cursor.fetchall():
            comp_name, avg_level = row
            avg_level = float(avg_level)
            competencies_text_content += f"{comp_name}: {avg_level:.2f}\n"

            if avg_level < 3:
                weak_zones_content += f"{comp_name}: {avg_level:.2f}\n"

                if "Презентация" in comp_name:
                    recommendations_content += "Вы почти не развиваете компетенцию 'Презентация результатов'. Рекомендуем выступить на студенческой конференции.\n\n"
                elif "Командная" in comp_name:
                    recommendations_content += "Низкий уровень командной работы. Участвуйте в групповых проектах.\n\n"
                elif "БД" in comp_name:
                    recommendations_content += "Слабая компетенция 'Работа с БД'. Пройдите дополнительный курс по базам данных.\n\n"

        if not recommendations_content:
            recommendations_content = "Все компетенции развиты хорошо. Продолжайте в том же духе!"

        self.competencies_text.config(state=tk.NORMAL)
        self.competencies_text.delete("1.0", tk.END)
        self.competencies_text.insert("1.0", competencies_text_content if competencies_text_content else "Нет данных")
        self.competencies_text.config(state=tk.DISABLED)

        self.weak_zones_text.config(state=tk.NORMAL)
        self.weak_zones_text.delete("1.0", tk.END)
        self.weak_zones_text.insert("1.0", weak_zones_content if weak_zones_content else "Слабых зон не обнаружено")
        self.weak_zones_text.config(state=tk.DISABLED)

        self.recommendations_text.config(state=tk.NORMAL)
        self.recommendations_text.delete("1.0", tk.END)
        self.recommendations_text.insert("1.0", recommendations_content)
        self.recommendations_text.config(state=tk.DISABLED)

    def check_achievements(self):
        self.cursor.execute("SELECT COUNT(*) FROM entries WHERE user_id = %s", (self.current_user_id,))
        total_entries = self.cursor.fetchone()[0]

        if total_entries == 1:
            self.unlock_achievement("Первый шаг", "Создана первая запись")

        self.cursor.execute("""
            SELECT COUNT(*) FROM entries 
            WHERE user_id = %s AND coauthors IS NOT NULL AND coauthors != ''
        """, (self.current_user_id,))
        entries_with_coauthors = self.cursor.fetchone()[0]

        if entries_with_coauthors >= 3:
            self.unlock_achievement("Командный игрок", "Три и более записи с соавторами")

        self.cursor.execute("""
            SELECT COUNT(DISTINCT type) FROM entries WHERE user_id = %s
        """, (self.current_user_id,))
        distinct_types = self.cursor.fetchone()[0]

        if distinct_types >= 3:
            self.unlock_achievement("Разносторонний", "Записи минимум трёх разных типов")

        self.cursor.execute("""
            SELECT EXTRACT(YEAR FROM date) as year, COUNT(*) as count
            FROM entries WHERE user_id = %s
            GROUP BY EXTRACT(YEAR FROM date)
            HAVING COUNT(*) >= 3
        """, (self.current_user_id,))

        if self.cursor.fetchone():
            self.unlock_achievement("Плодотворный год", "Три и более записи за один календарный год")

        self.cursor.execute("""
            SELECT SUM(LENGTH(description)) FROM entries WHERE user_id = %s
        """, (self.current_user_id,))
        total_chars = self.cursor.fetchone()[0] or 0

        if total_chars > 5000:
            self.unlock_achievement("Словобог", "Суммарный объём описаний превысил 5000 символов")

    def unlock_achievement(self, name, description):
        self.cursor.execute(
            "SELECT id FROM achievements WHERE name = %s AND user_id = %s",
            (name, self.current_user_id)
        )

        if not self.cursor.fetchone():
            self.cursor.execute(
                "INSERT INTO achievements (name, description, user_id, unlocked_date) VALUES (%s, %s, %s, %s)",
                (name, description, self.current_user_id, datetime.now().strftime("%Y-%m-%d"))
            )
            self.conn.commit()
            self.update_achievements()

    def update_achievements(self):
        self.cursor.execute(
            "SELECT name, description, unlocked_date FROM achievements WHERE user_id = %s ORDER BY unlocked_date DESC",
            (self.current_user_id,)
        )

        achievements_content = ""
        for row in self.cursor.fetchall():
            achievements_content += f"● {row[0]}\n"
            achievements_content += f"  {row[1]}\n"
            achievements_content += f"  Получено: {row[2]}\n\n"

        self.achievements_text.config(state=tk.NORMAL)
        self.achievements_text.delete("1.0", tk.END)
        self.achievements_text.insert("1.0", achievements_content if achievements_content else "Достижения отсутствуют")
        self.achievements_text.config(state=tk.DISABLED)

    def add_goal(self):
        description = self.goal_entry.get().strip()
        target_value = self.target_entry.get().strip()

        if not description:
            messagebox.showerror("Ошибка", "Введите описание цели")
            return

        try:
            target = int(target_value) if target_value.isdigit() else 1
        except ValueError:
            target = 1

        try:
            self.cursor.execute(
                "INSERT INTO goals (description, target_value, current_value, user_id) VALUES (%s, %s, %s, %s)",
                (description, target, 0, self.current_user_id)
            )
            self.conn.commit()

            self.goal_entry.delete(0, tk.END)
            self.target_entry.delete(0, tk.END)
            self.load_goals()

        except Exception as e:
            self.conn.rollback()
            messagebox.showerror("Ошибка базы данных", f"Не удалось добавить цель:\n{str(e)}")

    def load_goals(self):
        self.cursor.execute(
            "SELECT description, target_value, current_value FROM goals WHERE user_id = %s",
            (self.current_user_id,)
        )

        goals_content = ""
        for row in self.cursor.fetchall():
            description, target, current = row
            goals_content += f"Цель: {description}\n"
            goals_content += f"Прогресс: {current} из {target}\n"
            goals_content += f"Статус: {'Выполнено' if current >= target else 'В процессе'}\n\n"

        self.goals_text.config(state=tk.NORMAL)
        self.goals_text.delete("1.0", tk.END)
        self.goals_text.insert("1.0", goals_content if goals_content else "Цели не установлены")
        self.goals_text.config(state=tk.DISABLED)

    def export_to_word(self):
        doc = Document()
        doc.add_heading('Отчёт по портфолио', 0)

        doc.add_heading('1. Полный список записей', level=1)
        self.cursor.execute(
            "SELECT title, type, date, description, coauthors FROM entries WHERE user_id = %s ORDER BY date DESC",
            (self.current_user_id,)
        )

        for i, row in enumerate(self.cursor.fetchall(), 1):
            doc.add_heading(f'Запись {i}: {row[0]}', level=2)
            doc.add_paragraph(f'Тип: {row[1]}')
            doc.add_paragraph(f'Дата: {row[2]}')
            doc.add_paragraph(f'Описание: {row[3]}')
            doc.add_paragraph(f'Соавторы: {row[4] if row[4] else "Отсутствуют"}')

        doc.add_heading('2. Сводка по ключевым словам', level=1)
        self.cursor.execute("""
            SELECT k.keyword, COUNT(ek.entry_id) as count
            FROM keywords k
            JOIN entry_keywords ek ON k.id = ek.keyword_id
            JOIN entries e ON ek.entry_id = e.id
            WHERE e.user_id = %s
            GROUP BY k.keyword
            ORDER BY count DESC
        """, (self.current_user_id,))

        for row in self.cursor.fetchall():
            doc.add_paragraph(f'{row[0]} — {row[1]} записей')

        doc.add_heading('3. Сводка по соавторам', level=1)
        self.cursor.execute("""
            SELECT e.coauthors
            FROM entries e
            WHERE e.user_id = %s AND e.coauthors IS NOT NULL AND e.coauthors != ''
        """, (self.current_user_id,))

        coauthors_dict = {}
        for row in self.cursor.fetchall():
            coauthors = [c.strip() for c in row[0].split(",") if c.strip()]
            for coauthor in coauthors:
                coauthors_dict[coauthor] = coauthors_dict.get(coauthor, 0) + 1

        for coauthor, count in sorted(coauthors_dict.items(), key=lambda x: x[1], reverse=True):
            doc.add_paragraph(f'{coauthor} — {count} работ')

        doc.add_heading('4. Профиль компетенций', level=1)
        self.cursor.execute("""
            SELECT c.name, CAST(AVG(ec.level) AS DECIMAL(10,2)) as avg_level
            FROM competencies c
            LEFT JOIN entry_competencies ec ON c.id = ec.competency_id
            LEFT JOIN entries e ON ec.entry_id = e.id AND e.user_id = %s
            GROUP BY c.id, c.name
            HAVING AVG(ec.level) IS NOT NULL
            ORDER BY AVG(ec.level) DESC
        """, (self.current_user_id,))

        for row in self.cursor.fetchall():
            doc.add_paragraph(f'{row[0]}: {float(row[1]):.2f}')

        doc.add_heading('5. Персонализированные рекомендации', level=1)
        self.cursor.execute("""
            SELECT c.name, CAST(AVG(ec.level) AS DECIMAL(10,2)) as avg_level
            FROM competencies c
            LEFT JOIN entry_competencies ec ON c.id = ec.competency_id
            LEFT JOIN entries e ON ec.entry_id = e.id AND e.user_id = %s
            GROUP BY c.id, c.name
            HAVING AVG(ec.level) IS NOT NULL AND AVG(ec.level) < 3
        """, (self.current_user_id,))

        weak_comps = self.cursor.fetchall()
        if weak_comps:
            doc.add_paragraph('Рекомендации по развитию слабых зон:')
            for comp_name, avg_level in weak_comps:
                avg_level = float(avg_level)
                if "Презентация" in comp_name:
                    doc.add_paragraph(
                        f'- Для компетенции "{comp_name}" (уровень: {avg_level:.2f}): выступите на студенческой конференции')
                elif "Командная" in comp_name:
                    doc.add_paragraph(
                        f'- Для компетенции "{comp_name}" (уровень: {avg_level:.2f}): участвуйте в групповых проектах')
                elif "БД" in comp_name:
                    doc.add_paragraph(
                        f'- Для компетенции "{comp_name}" (уровень: {avg_level:.2f}): пройдите курс по базам данных')
        else:
            doc.add_paragraph('Все компетенции развиты хорошо. Продолжайте в том же духе!')

        doc.add_heading('6. Список полученных достижений', level=1)
        self.cursor.execute(
            "SELECT name, description, unlocked_date FROM achievements WHERE user_id = %s",
            (self.current_user_id,)
        )

        for row in self.cursor.fetchall():
            doc.add_paragraph(f'● {row[0]}: {row[1]} (получено {row[2]})')

        filename = f"portfolio_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        messagebox.showinfo("Экспорт завершён", f"Отчёт сохранён в файл: {filename}")

    def run(self):
        self.root.mainloop()

    def __del__(self):
        if hasattr(self, 'cursor'):
            self.cursor.close()
        if hasattr(self, 'conn'):
            self.conn.close()


if __name__ == "__main__":
    root = tk.Tk()
    app = PortfolioApp(root)
    app.run()
