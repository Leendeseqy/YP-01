import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import sqlite3
import json
import os
from datetime import datetime
from docx import Document
import pandas as pd
from tkcalendar import DateEntry
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import hashlib
import re
import schedule
import threading
import time
import numpy as np
from collections import Counter
import shutil


class AchievementTracker:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("📚 Журнал личных учебных достижений")
        self.root.geometry("1000x700")
        self.root.configure(bg='#f0f0f0')

        # Загрузка типов из JSON
        self.achievement_types = self.load_types()

        # Уровни достижений
        self.levels = ["локальный", "региональный", "национальный", "международный"]

        # Инициализация базы данных
        self.init_db()

        # Создание интерфейса
        self.create_ui()

        # Запуск планировщика уведомлений
        self.start_notification_scheduler()

    def load_types(self):
        """Загрузка типов достижений из JSON файла"""
        try:
            with open("types.json", "r", encoding="utf-8") as f:
                types = json.load(f)
                if not isinstance(types, list):
                    return ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]
                return types
        except:
            return ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]

    def init_db(self):
        """Инициализация базы данных"""
        self.conn = sqlite3.connect("достижения.db", check_same_thread=False)
        self.cursor = self.conn.cursor()
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS достижения (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                название TEXT NOT NULL,
                дата TEXT NOT NULL,
                тип TEXT NOT NULL,
                уровень TEXT NOT NULL,
                описание TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Создание таблицы для статистики
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS статистика (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                тип TEXT NOT NULL,
                количество INTEGER DEFAULT 0,
                месяц TEXT NOT NULL
            )
        """)
        self.conn.commit()

    def create_ui(self):
        """Создание пользовательского интерфейса"""
        # Стилизация
        style = ttk.Style()
        style.theme_use('clam')

        # Создание Notebook (вкладок)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Вкладки
        self.tab_add = ttk.Frame(self.notebook)
        self.tab_list = ttk.Frame(self.notebook)
        self.tab_stats = ttk.Frame(self.notebook)
        self.tab_search = ttk.Frame(self.notebook)

        self.notebook.add(self.tab_add, text="➕ Добавить")
        self.notebook.add(self.tab_list, text="📋 Мои достижения")
        self.notebook.add(self.tab_stats, text="📊 Статистика")
        self.notebook.add(self.tab_search, text="🔍 Поиск")

        # Создание форм на каждой вкладке
        self.create_add_form()
        self.create_list_form()
        self.create_stats_form()
        self.create_search_form()

    def create_add_form(self):
        """Создание формы добавления достижений"""
        # Фрейм для формы
        form_frame = ttk.LabelFrame(self.tab_add, text="Добавить новое достижение", padding=20)
        form_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Название
        ttk.Label(form_frame, text="Название:*", font=('Arial', 11, 'bold')).grid(row=0, column=0, sticky='w',
                                                                                  pady=(0, 5))
        self.name_entry = ttk.Entry(form_frame, width=50, font=('Arial', 11))
        self.name_entry.grid(row=0, column=1, padx=(10, 0), pady=(0, 15))

        # Дата с календарем
        ttk.Label(form_frame, text="Дата:*", font=('Arial', 11, 'bold')).grid(row=1, column=0, sticky='w', pady=(0, 5))
        self.date_entry = DateEntry(form_frame, width=47, background='darkblue',
                                    foreground='white', borderwidth=2, date_pattern='yyyy-mm-dd',
                                    font=('Arial', 11))
        self.date_entry.grid(row=1, column=1, padx=(10, 0), pady=(0, 15))

        # Тип достижения
        ttk.Label(form_frame, text="Тип:*", font=('Arial', 11, 'bold')).grid(row=2, column=0, sticky='w', pady=(0, 5))
        self.type_combo = ttk.Combobox(form_frame, values=self.achievement_types,
                                       state="readonly", font=('Arial', 11), width=47)
        self.type_combo.grid(row=2, column=1, padx=(10, 0), pady=(0, 15))
        self.type_combo.current(0)

        # Уровень
        ttk.Label(form_frame, text="Уровень:*", font=('Arial', 11, 'bold')).grid(row=3, column=0, sticky='w',
                                                                                 pady=(0, 5))
        self.level_combo = ttk.Combobox(form_frame, values=self.levels,
                                        state="readonly", font=('Arial', 11), width=47)
        self.level_combo.grid(row=3, column=1, padx=(10, 0), pady=(0, 15))
        self.level_combo.current(0)

        # Описание
        ttk.Label(form_frame, text="Описание:", font=('Arial', 11, 'bold')).grid(row=4, column=0, sticky='nw',
                                                                                 pady=(0, 5))
        self.desc_text = tk.Text(form_frame, height=6, width=50, font=('Arial', 11), wrap='word')
        self.desc_text.grid(row=4, column=1, padx=(10, 0), pady=(0, 20))

        # Кнопки
        button_frame = ttk.Frame(form_frame)
        button_frame.grid(row=5, column=0, columnspan=2, pady=20)

        ttk.Button(button_frame, text="💾 Сохранить", command=self.save_achievement,
                   style='Accent.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text="🧹 Очистить", command=self.clear_form).pack(side='left', padx=5)

        # Обязательные поля
        ttk.Label(form_frame, text="* - обязательные поля", font=('Arial', 9, 'italic'),
                  foreground='red').grid(row=6, column=0, columnspan=2, pady=(10, 0))

        # Стиль для акцентной кнопки
        style = ttk.Style()
        style.configure('Accent.TButton', background='#4CAF50', foreground='white', font=('Arial', 11, 'bold'))

    def create_list_form(self):
        """Создание формы списка достижений - ИСПРАВЛЕННАЯ ВЕРСИЯ"""
        # Панель управления
        control_frame = ttk.Frame(self.tab_list)
        control_frame.pack(fill='x', padx=20, pady=(20, 10))

        # Кнопки экспорта слева
        export_frame = ttk.Frame(control_frame)
        export_frame.pack(side='left')

        ttk.Button(export_frame, text="📄 Word", command=self.export_to_word, width=10).pack(side='left', padx=2)
        ttk.Button(export_frame, text="📊 Excel", command=self.export_to_excel, width=10).pack(side='left', padx=2)
        ttk.Button(export_frame, text="📈 PDF", command=self.export_to_pdf, width=10).pack(side='left', padx=2)

        # Кнопки действий справа
        action_frame = ttk.Frame(control_frame)
        action_frame.pack(side='right')

        ttk.Button(action_frame, text="🗑️ Удалить", command=self.delete_record, width=10).pack(side='left', padx=2)
        ttk.Button(action_frame, text="👁️ Просмотр", command=self.view_details, width=10).pack(side='left', padx=2)
        ttk.Button(action_frame, text="✏️ Редактировать", command=self.edit_record, width=12).pack(side='left', padx=2)
        ttk.Button(action_frame, text="🔄 Обновить", command=self.refresh_list, width=10).pack(side='left', padx=2)

        # Фрейм для дерева с прокруткой
        tree_frame = ttk.Frame(self.tab_list)
        tree_frame.pack(fill='both', expand=True, padx=20, pady=(0, 10))

        # Дерево для отображения записей
        columns = ('Дата', 'Название', 'Тип', 'Уровень', 'Описание')
        self.tree = ttk.Treeview(tree_frame, columns=columns, show='headings', height=20)

        # Настройка колонок с правильными ширинами
        self.tree.heading('Дата', text='Дата')
        self.tree.heading('Название', text='Название')
        self.tree.heading('Тип', text='Тип')
        self.tree.heading('Уровень', text='Уровень')
        self.tree.heading('Описание', text='Описание')

        self.tree.column('Дата', width=100, minwidth=80)
        self.tree.column('Название', width=200, minwidth=150)
        self.tree.column('Тип', width=120, minwidth=80)
        self.tree.column('Уровень', width=120, minwidth=80)
        self.tree.column('Описание', width=300, minwidth=150)

        # Вертикальная прокрутка
        v_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=v_scrollbar.set)

        # Горизонтальная прокрутка
        h_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(xscrollcommand=h_scrollbar.set)

        # Размещение элементов
        self.tree.grid(row=0, column=0, sticky='nsew')
        v_scrollbar.grid(row=0, column=1, sticky='ns')
        h_scrollbar.grid(row=1, column=0, sticky='ew')

        # Настройка расширения
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)

        # Стилизация дерева
        style = ttk.Style()
        style.configure("Treeview",
                        background="#ffffff",
                        foreground="#000000",
                        rowheight=25,
                        fieldbackground="#ffffff")

        style.configure("Treeview.Heading",
                        background="#4CAF50",
                        foreground="white",
                        font=('Arial', 10, 'bold'))

        style.map('Treeview.Heading',
                  background=[('active', '#45a049')])

        # Контекстное меню
        self.tree.bind('<Button-3>', self.show_context_menu)
        self.tree.bind('<Double-Button-1>', lambda e: self.view_details())

        self.context_menu = tk.Menu(self.root, tearoff=0)
        self.context_menu.add_command(label="👁️ Просмотр деталей", command=self.view_details)
        self.context_menu.add_command(label="✏️ Редактировать", command=self.edit_record)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="🗑️ Удалить", command=self.delete_record)

        # Панель статуса внизу
        self.status_frame = ttk.Frame(self.tab_list)
        self.status_frame.pack(fill='x', padx=20, pady=(0, 10))

        self.status_label = ttk.Label(self.status_frame, text="Готово", font=('Arial', 9))
        self.status_label.pack(side='left')

        self.count_label = ttk.Label(self.status_frame, text="Записей: 0", font=('Arial', 9, 'bold'))
        self.count_label.pack(side='right')

        # Привязка событий
        self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)

        # Загрузка данных
        self.refresh_list()

    def create_stats_form(self):
        """Создание формы статистики"""
        # Фрейм для графиков
        stats_frame = ttk.Frame(self.tab_stats)
        stats_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Кнопки обновления статистики
        button_frame = ttk.Frame(stats_frame)
        button_frame.pack(fill='x', pady=(0, 20))

        ttk.Button(button_frame, text="🔄 Обновить статистику",
                   command=self.update_stats).pack(side='left', padx=5)
        ttk.Button(button_frame, text="📈 Показать графики",
                   command=self.show_charts).pack(side='left', padx=5)

        # Фрейм для отображения статистики
        self.stats_text = tk.Text(stats_frame, height=20, width=80, font=('Arial', 11))
        self.stats_text.pack(fill='both', expand=True)

        # Инициализация статистики
        self.update_stats()

    def create_search_form(self):
        """Создание формы поиска"""
        # Панель поиска
        search_frame = ttk.LabelFrame(self.tab_search, text="Параметры поиска", padding=15)
        search_frame.pack(fill='x', padx=20, pady=20)

        # Поиск по названию
        ttk.Label(search_frame, text="Название:", font=('Arial', 10)).grid(row=0, column=0, sticky='w', pady=5)
        self.search_name = ttk.Entry(search_frame, width=30)
        self.search_name.grid(row=0, column=1, padx=(10, 20), pady=5)

        # Фильтр по типу
        ttk.Label(search_frame, text="Тип:", font=('Arial', 10)).grid(row=0, column=2, sticky='w', pady=5)
        self.search_type = ttk.Combobox(search_frame, values=["Все"] + self.achievement_types,
                                        state="readonly", width=20)
        self.search_type.grid(row=0, column=3, padx=(10, 0), pady=5)
        self.search_type.current(0)

        # Фильтр по уровню
        ttk.Label(search_frame, text="Уровень:", font=('Arial', 10)).grid(row=1, column=0, sticky='w', pady=5)
        self.search_level = ttk.Combobox(search_frame, values=["Все"] + self.levels,
                                         state="readonly", width=20)
        self.search_level.grid(row=1, column=1, padx=(10, 20), pady=5)
        self.search_level.current(0)

        # Диапазон дат
        ttk.Label(search_frame, text="С даты:", font=('Arial', 10)).grid(row=1, column=2, sticky='w', pady=5)
        self.date_from = DateEntry(search_frame, width=18, date_pattern='yyyy-mm-dd')
        self.date_from.grid(row=1, column=3, padx=(10, 0), pady=5)

        ttk.Label(search_frame, text="По дату:", font=('Arial', 10)).grid(row=2, column=2, sticky='w', pady=5)
        self.date_to = DateEntry(search_frame, width=18, date_pattern='yyyy-mm-dd')
        self.date_to.grid(row=2, column=3, padx=(10, 0), pady=5)

        # Кнопки поиска
        button_frame = ttk.Frame(search_frame)
        button_frame.grid(row=3, column=0, columnspan=4, pady=15)

        ttk.Button(button_frame, text="🔍 Искать", command=self.perform_search,
                   style='Accent.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text="🧹 Сбросить", command=self.reset_search).pack(side='left', padx=5)

        # Результаты поиска
        result_frame = ttk.LabelFrame(self.tab_search, text="Результаты поиска", padding=10)
        result_frame.pack(fill='both', expand=True, padx=20, pady=(0, 20))

        columns = ('Дата', 'Название', 'Тип', 'Уровень', 'Описание')
        self.search_tree = ttk.Treeview(result_frame, columns=columns, show='headings', height=15)

        for col in columns:
            self.search_tree.heading(col, text=col)
            self.search_tree.column(col, width=120)

        scrollbar = ttk.Scrollbar(result_frame, orient="vertical", command=self.search_tree.yview)
        self.search_tree.configure(yscrollcommand=scrollbar.set)

        self.search_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

    def validate_input(self, name, date_str):
        """Валидация введенных данных"""
        errors = []

        if not name or len(name.strip()) < 3:
            errors.append("Название должно содержать минимум 3 символа")

        # Проверка даты с помощью регулярного выражения
        date_pattern = r'^\d{4}-\d{2}-\d{2}$'
        if not re.match(date_pattern, date_str):
            errors.append("Дата должна быть в формате ГГГГ-ММ-ДД")
        else:
            try:
                datetime.strptime(date_str, '%Y-%m-%d')
            except ValueError:
                errors.append("Некорректная дата")

        return errors

    def save_achievement(self):
        """Сохранение достижения в базу данных"""
        name = self.name_entry.get().strip()
        date_str = self.date_entry.get()
        typ = self.type_combo.get()
        level = self.level_combo.get()
        desc = self.desc_text.get("1.0", "end-1c").strip()

        # Валидация
        errors = self.validate_input(name, date_str)
        if errors:
            messagebox.showerror("Ошибка валидации", "\n".join(errors))
            return

        try:
            # Хеширование для безопасности (демонстрация)
            name_hash = hashlib.md5(name.encode()).hexdigest()[:8]

            self.cursor.execute("""
                INSERT INTO достижения (название, дата, тип, уровень, описание)
                VALUES (?, ?, ?, ?, ?)
            """, (name, date_str, typ, level, desc))

            self.conn.commit()

            # Обновление статистики
            self.update_statistics(typ)

            messagebox.showinfo("Успех", f"Достижение '{name}' успешно сохранено!\nID: {name_hash}")
            self.clear_form()
            self.refresh_list()

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить данные: {str(e)}")

    def clear_form(self):
        """Очистка формы ввода"""
        self.name_entry.delete(0, tk.END)
        self.date_entry.set_date(datetime.now())
        self.type_combo.current(0)
        self.level_combo.current(0)
        self.desc_text.delete("1.0", tk.END)

    def load_records(self, with_description=False):
        """Загрузка записей из базы данных"""
        try:
            if with_description:
                self.cursor.execute("""
                    SELECT id, дата, название, тип, уровень, описание 
                    FROM достижения 
                    ORDER BY дата DESC
                """)
                return self.cursor.fetchall()
            else:
                # ВАЖНО: Загружаем ВСЕ поля включая описание
                self.cursor.execute("""
                    SELECT дата, название, тип, уровень, описание 
                    FROM достижения 
                    ORDER BY дата DESC
                """)
                return self.cursor.fetchall()
        except Exception as e:
            print(f"Ошибка загрузки записей: {e}")
            return []

    def refresh_list(self):
        """Обновление списка достижений - ИСПРАВЛЕННАЯ ВЕРСИЯ"""
        try:
            # Очистка дерева
            for item in self.tree.get_children():
                self.tree.delete(item)

            # Загрузка записей с описанием
            self.cursor.execute("""
                SELECT дата, название, тип, уровень, описание 
                FROM достижения 
                ORDER BY дата DESC
            """)
            records = self.cursor.fetchall()

            # Отображение записей с обработкой None значений
            for record in records:
                # Преобразуем None в пустые строки для отображения
                processed_record = []
                for value in record:
                    if value is None:
                        processed_record.append("")
                    else:
                        # Обрезаем слишком длинные описания для лучшего отображения
                        if isinstance(value, str) and len(value) > 100:
                            processed_record.append(value[:97] + "...")
                        else:
                            processed_record.append(str(value))

                # ВСТАВЛЯЕМ все 5 значений: дата, название, тип, уровень, описание
                self.tree.insert('', 'end', values=processed_record)

            # Обновляем статус в заголовке вкладки
            tab_index = self.notebook.index(self.tab_list)
            self.notebook.tab(tab_index, text=f"📋 Мои достижения ({len(records)})")

            # Обновляем счетчик в статусной панели
            self.count_label.config(text=f"Записей: {len(records)}")

            self.status_label.config(text=f"Загружено {len(records)} записей")

            # Показываем предупреждение если нет записей
            if len(records) == 0:
                self.status_label.config(text="Нет данных. Добавьте достижения на вкладке '➕ Добавить'")

        except Exception as e:
            error_msg = f"Ошибка обновления списка: {str(e)}"
            print(error_msg)
            self.status_label.config(text=error_msg)
            messagebox.showerror("Ошибка", error_msg)

    def on_tree_select(self, event):
        """Обработка выбора элемента в дереве"""
        selection = self.tree.selection()
        if selection:
            item = self.tree.item(selection[0])
            self.status_label.config(text=f"Выбрано: {item['values'][1]}")
        else:
            self.status_label.config(text="Готово")

    def show_context_menu(self, event):
        """Показ контекстного меню"""
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)

    def edit_record(self):
        """Редактирование выбранной записи"""
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите запись для редактирования")
            return

        # Получение данных записи
        item = self.tree.item(selection[0])
        values = item['values']

        # Открытие окна редактирования
        self.open_edit_window(values)

    def delete_record(self):
        """Удаление выбранной записи"""
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите запись для удаления")
            return

        item = self.tree.item(selection[0])
        achievement_name = item['values'][1]

        if messagebox.askyesno("Подтверждение", f"Удалить достижение '{achievement_name}'?"):
            try:
                date_str = item['values'][0]
                achievement_type = item['values'][2]

                # Удаляем запись
                self.cursor.execute("""
                    DELETE FROM достижения 
                    WHERE дата = ? AND название = ? AND тип = ?
                """, (date_str, achievement_name, achievement_type))

                self.conn.commit()

                # Обновляем статистику
                self.update_statistics_after_delete(achievement_type)

                # Обновляем список
                self.refresh_list()

                messagebox.showinfo("Успех", "Запись успешно удалена")

            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить запись: {str(e)}")

    def view_details(self):
        """Просмотр деталей записи"""
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите запись для просмотра")
            return

        item = self.tree.item(selection[0])
        values = item['values']

        # Создание окна с деталями
        detail_window = tk.Toplevel(self.root)
        detail_window.title("Детали достижения")
        detail_window.geometry("500x400")
        detail_window.configure(bg='#f0f0f0')

        # Заголовок
        header_frame = ttk.Frame(detail_window)
        header_frame.pack(fill='x', padx=20, pady=(20, 10))

        ttk.Label(header_frame, text=values[1], font=('Arial', 14, 'bold'),
                  wraplength=400, justify='center').pack()

        # Информация
        info_frame = ttk.LabelFrame(detail_window, text="Информация", padding=15)
        info_frame.pack(fill='both', expand=True, padx=20, pady=10)

        # Дата
        date_frame = ttk.Frame(info_frame)
        date_frame.pack(fill='x', pady=5)
        ttk.Label(date_frame, text="Дата:", font=('Arial', 11, 'bold'), width=10).pack(side='left')
        ttk.Label(date_frame, text=values[0], font=('Arial', 11)).pack(side='left')

        # Тип
        type_frame = ttk.Frame(info_frame)
        type_frame.pack(fill='x', pady=5)
        ttk.Label(type_frame, text="Тип:", font=('Arial', 11, 'bold'), width=10).pack(side='left')
        ttk.Label(type_frame, text=values[2], font=('Arial', 11)).pack(side='left')

        # Уровень
        level_frame = ttk.Frame(info_frame)
        level_frame.pack(fill='x', pady=5)
        ttk.Label(level_frame, text="Уровень:", font=('Arial', 11, 'bold'), width=10).pack(side='left')
        ttk.Label(level_frame, text=values[3], font=('Arial', 11)).pack(side='left')

        # Описание
        desc_frame = ttk.LabelFrame(info_frame, text="Описание", padding=10)
        desc_frame.pack(fill='both', expand=True, pady=(15, 0))

        desc_text = tk.Text(desc_frame, height=8, wrap='word', font=('Arial', 11))
        desc_text.pack(fill='both', expand=True, side='left')

        # Заполняем описание
        description = values[4] if len(values) > 4 else ""
        if not description:
            description = "Описание отсутствует"

        desc_text.insert('1.0', description)
        desc_text.config(state='disabled')

        # Скроллбар
        scrollbar = ttk.Scrollbar(desc_frame, command=desc_text.yview)
        desc_text.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side='right', fill='y')

        # Кнопка закрытия
        ttk.Button(detail_window, text="Закрыть", command=detail_window.destroy,
                   style='Accent.TButton').pack(pady=20)

    def open_edit_window(self, values):
        """Открытие окна редактирования"""
        edit_window = tk.Toplevel(self.root)
        edit_window.title("Редактирование достижения")
        edit_window.geometry("500x500")

        # Загрузка полной записи по уникальным данным
        self.cursor.execute("""
            SELECT id, описание FROM достижения 
            WHERE дата = ? AND название = ? AND тип = ? AND уровень = ?
        """, values[:4])

        record = self.cursor.fetchone()
        if not record:
            messagebox.showerror("Ошибка", "Запись не найдена в базе данных")
            return

        record_id = record[0]
        current_desc = record[1] if record[1] else ""

        # Форма редактирования
        form_frame = ttk.Frame(edit_window, padding=20)
        form_frame.pack(fill='both', expand=True)

        ttk.Label(form_frame, text="Название:*", font=('Arial', 11, 'bold')).grid(row=0, column=0, sticky='w', pady=5)
        name_entry = ttk.Entry(form_frame, width=40, font=('Arial', 11))
        name_entry.insert(0, values[1])
        name_entry.grid(row=0, column=1, padx=10, pady=5)

        ttk.Label(form_frame, text="Дата:*", font=('Arial', 11, 'bold')).grid(row=1, column=0, sticky='w', pady=5)
        date_entry = ttk.Entry(form_frame, width=40, font=('Arial', 11))
        date_entry.insert(0, values[0])
        date_entry.grid(row=1, column=1, padx=10, pady=5)

        ttk.Label(form_frame, text="Тип:*", font=('Arial', 11, 'bold')).grid(row=2, column=0, sticky='w', pady=5)
        type_combo = ttk.Combobox(form_frame, values=self.achievement_types, state="readonly", width=37)
        type_combo.set(values[2])
        type_combo.grid(row=2, column=1, padx=10, pady=5)

        ttk.Label(form_frame, text="Уровень:*", font=('Arial', 11, 'bold')).grid(row=3, column=0, sticky='w', pady=5)
        level_combo = ttk.Combobox(form_frame, values=self.levels, state="readonly", width=37)
        level_combo.set(values[3])
        level_combo.grid(row=3, column=1, padx=10, pady=5)

        ttk.Label(form_frame, text="Описание:", font=('Arial', 11, 'bold')).grid(row=4, column=0, sticky='nw', pady=5)
        desc_text = tk.Text(form_frame, height=8, width=40, wrap='word', font=('Arial', 11))
        desc_text.insert('1.0', current_desc)
        desc_text.grid(row=4, column=1, padx=10, pady=5)

        def save_changes():
            new_name = name_entry.get().strip()
            new_date = date_entry.get().strip()
            new_type = type_combo.get()
            new_level = level_combo.get()
            new_desc = desc_text.get("1.0", "end-1c").strip()

            errors = self.validate_input(new_name, new_date)
            if errors:
                messagebox.showerror("Ошибка валидации", "\n".join(errors))
                return

            try:
                self.cursor.execute("""
                    UPDATE достижения 
                    SET название = ?, дата = ?, тип = ?, уровень = ?, описание = ?
                    WHERE id = ?
                """, (new_name, new_date, new_type, new_level, new_desc, record_id))

                self.conn.commit()
                messagebox.showinfo("Успех", "Изменения сохранены")
                edit_window.destroy()
                self.refresh_list()

            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить изменения: {str(e)}")

        button_frame = ttk.Frame(form_frame)
        button_frame.grid(row=5, column=0, columnspan=2, pady=20)

        ttk.Button(button_frame, text="💾 Сохранить", command=save_changes,
                   style='Accent.TButton').pack(side='left', padx=10)
        ttk.Button(button_frame, text="❌ Отмена", command=edit_window.destroy).pack(side='left', padx=10)

    def export_to_word(self):
        """Экспорт в Word документ"""
        try:
            doc = Document()

            # Заголовок
            doc.add_heading('Личные учебные достижения', 0)
            doc.add_paragraph(f'Отчет создан: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            doc.add_paragraph()

            # Загрузка данных
            records = self.load_records(with_description=True)

            if not records:
                doc.add_paragraph("Нет данных для экспорта")
            else:
                # Таблица с данными
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'

                # Заголовки таблицы
                hdr_cells = table.rows[0].cells
                headers = ['Дата', 'Название', 'Тип', 'Уровень', 'Описание']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

                # Данные
                for record in records:
                    row_cells = table.add_row().cells
                    for i in range(5):
                        value = str(record[i + 1] if i < 5 else "")
                        row_cells[i].text = value

            # Сохранение файла
            filename = f"достижения_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)

            messagebox.showinfo("Экспорт завершен",
                                f"Документ успешно сохранен как:\n{filename}")

        except Exception as e:
            messagebox.showerror("Ошибка экспорта", f"Не удалось создать Word документ: {str(e)}")

    def export_to_excel(self):
        """Экспорт в Excel"""
        try:
            records = self.load_records(with_description=True)

            if not records:
                messagebox.showwarning("Нет данных", "Нет данных для экспорта")
                return

            # Создание DataFrame
            df = pd.DataFrame(records, columns=['ID', 'Дата', 'Название', 'Тип', 'Уровень', 'Описание'])
            df = df.drop('ID', axis=1)  # Удаляем ID

            # Сохранение в Excel
            filename = f"достижения_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            df.to_excel(filename, index=False, engine='openpyxl')

            messagebox.showinfo("Экспорт завершен",
                                f"Excel файл успешно сохранен как:\n{filename}")

        except Exception as e:
            messagebox.showerror("Ошибка экспорта", f"Не удалось создать Excel файл: {str(e)}")

    def export_to_pdf(self):
        """Экспорт в PDF (упрощенный через сохранение текста)"""
        try:
            records = self.load_records(with_description=True)

            if not records:
                messagebox.showwarning("Нет данных", "Нет данных для экспорта")
                return

            # Создание текстового файла как упрощенный PDF
            filename = f"достижения_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

            with open(filename, 'w', encoding='utf-8') as f:
                f.write("=" * 60 + "\n")
                f.write("ЛИЧНЫЕ УЧЕБНЫЕ ДОСТИЖЕНИЯ\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"Отчет создан: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

                for record in records:
                    f.write(f"Дата: {record[1]}\n")
                    f.write(f"Название: {record[2]}\n")
                    f.write(f"Тип: {record[3]}\n")
                    f.write(f"Уровень: {record[4]}\n")
                    desc = record[5] if record[5] else "Нет описания"
                    f.write(f"Описание: {desc}\n")
                    f.write("-" * 60 + "\n")

            messagebox.showinfo("Экспорт завершен",
                                f"Текстовый отчет сохранен как:\n{filename}\n\nМожно конвертировать в PDF с помощью принтера.")

        except Exception as e:
            messagebox.showerror("Ошибка экспорта", f"Не удалось создать отчет: {str(e)}")

    def update_stats(self):
        """Обновление статистики"""
        try:
            # Включаем редактирование
            self.stats_text.config(state='normal')

            # Очистка текстового поля
            self.stats_text.delete('1.0', tk.END)

            # Получение ОБЩЕГО количества достижений
            self.cursor.execute("SELECT COUNT(*) FROM достижения")
            total_count_result = self.cursor.fetchone()
            total_count = total_count_result[0] if total_count_result else 0

            # Статистика по типам
            self.cursor.execute("""
                SELECT тип, COUNT(*) as count 
                FROM достижения 
                GROUP BY тип 
                ORDER BY count DESC
            """)
            type_stats = self.cursor.fetchall()

            # Статистика по уровням
            self.cursor.execute("""
                SELECT уровень, COUNT(*) as count 
                FROM достижения 
                GROUP BY уровень 
                ORDER BY count DESC
            """)
            level_stats = self.cursor.fetchall()

            # Статистика по месяцам
            self.cursor.execute("""
                SELECT strftime('%Y-%m', дата) as месяц, COUNT(*) as count
                FROM достижения 
                WHERE дата IS NOT NULL AND дата != ''
                GROUP BY месяц 
                ORDER BY месяц DESC
                LIMIT 12
            """)
            month_stats = self.cursor.fetchall()

            # Форматирование вывода
            self.stats_text.insert('1.0', "📊 СТАТИСТИКА ДОСТИЖЕНИЙ\n")
            self.stats_text.insert('end', "=" * 50 + "\n\n")

            self.stats_text.insert('end', f"📈 Всего достижений: {total_count}\n\n")

            if total_count > 0:
                self.stats_text.insert('end', "🏆 По типам:\n")
                for typ, count in type_stats:
                    percentage = (count / total_count * 100) if total_count > 0 else 0
                    bar_length = int(percentage / 2)
                    bar = "█" * bar_length if bar_length > 0 else ""
                    self.stats_text.insert('end', f"  {typ}: {count} ({percentage:.1f}%) {bar}\n")

                self.stats_text.insert('end', "\n📊 По уровням:\n")
                for level, count in level_stats:
                    percentage = (count / total_count * 100) if total_count > 0 else 0
                    self.stats_text.insert('end', f"  {level}: {count} ({percentage:.1f}%)\n")

                if month_stats:
                    self.stats_text.insert('end', "\n📅 По месяцам (последние 12):\n")
                    for month, count in month_stats:
                        self.stats_text.insert('end', f"  {month}: {count}\n")

                # Последнее достижение
                self.cursor.execute("""
                    SELECT название, дата 
                    FROM достижения 
                    WHERE дата IS NOT NULL AND дата != ''
                    ORDER BY дата DESC 
                    LIMIT 1
                """)
                last_record = self.cursor.fetchone()

                if last_record:
                    self.stats_text.insert('end', f"\n⏰ Последнее достижение:\n")
                    self.stats_text.insert('end', f"  {last_record[0]} ({last_record[1]})\n")
            else:
                self.stats_text.insert('end', "📭 Нет данных для отображения статистики\n")
                self.stats_text.insert('end', "Добавьте достижения на вкладке '➕ Добавить'\n")

            # Делаем текст read-only
            self.stats_text.config(state='disabled')

        except Exception as e:
            self.stats_text.insert('end', f"⚠️ Ошибка при загрузке статистики: {str(e)}\n")

    def update_statistics(self, achievement_type):
        """Обновление статистики в базе данных"""
        current_month = datetime.now().strftime('%Y-%m')

        # Проверяем, есть ли запись за этот месяц
        self.cursor.execute("""
            SELECT количество FROM статистика 
            WHERE тип = ? AND месяц = ?
        """, (achievement_type, current_month))

        result = self.cursor.fetchone()

        if result:
            # Обновляем существующую запись
            new_count = result[0] + 1
            self.cursor.execute("""
                UPDATE статистика 
                SET количество = ? 
                WHERE тип = ? AND месяц = ?
            """, (new_count, achievement_type, current_month))
        else:
            # Создаем новую запись
            self.cursor.execute("""
                INSERT INTO статистика (тип, количество, месяц)
                VALUES (?, 1, ?)
            """, (achievement_type, current_month))

        self.conn.commit()

    def update_statistics_after_delete(self, achievement_type):
        """Обновление статистики после удаления"""
        current_month = datetime.now().strftime('%Y-%m')

        # Получаем текущее количество
        self.cursor.execute("""
            SELECT количество FROM статистика 
            WHERE тип = ? AND месяц = ?
        """, (achievement_type, current_month))

        result = self.cursor.fetchone()

        if result:
            new_count = max(0, result[0] - 1)  # Не меньше 0

            if new_count > 0:
                self.cursor.execute("""
                    UPDATE статистика 
                    SET количество = ? 
                    WHERE тип = ? AND месяц = ?
                """, (new_count, achievement_type, current_month))
            else:
                # Удаляем запись если количество 0
                self.cursor.execute("""
                    DELETE FROM статистика 
                    WHERE тип = ? AND месяц = ?
                """, (achievement_type, current_month))

            self.conn.commit()

    def show_charts(self):
        """Отображение графиков статистики"""
        try:
            # Получение данных для графиков
            self.cursor.execute("SELECT тип, COUNT(*) FROM достижения GROUP BY тип")
            type_data = self.cursor.fetchall()

            self.cursor.execute("SELECT уровень, COUNT(*) FROM достижения GROUP BY уровень")
            level_data = self.cursor.fetchall()

            if not type_data:
                messagebox.showinfo("Нет данных", "Недостаточно данных для построения графиков")
                return

            # Создание окна с графиками
            chart_window = tk.Toplevel(self.root)
            chart_window.title("Графики статистики")
            chart_window.geometry("900x600")

            # Создание фигуры с двумя графиками
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

            # Первый график - по типам
            types = [item[0] for item in type_data]
            counts = [item[1] for item in type_data]

            # Цвета для графиков
            colors1 = plt.cm.Set3(np.linspace(0, 1, len(types)))
            bars1 = ax1.bar(types, counts, color=colors1)
            ax1.set_title('Распределение по типам', fontsize=14, fontweight='bold')
            ax1.set_xlabel('Тип достижения', fontsize=12)
            ax1.set_ylabel('Количество', fontsize=12)
            ax1.tick_params(axis='x', rotation=45, labelsize=10)
            ax1.grid(axis='y', alpha=0.3)

            # Добавление значений на столбцы
            for bar in bars1:
                height = bar.get_height()
                ax1.text(bar.get_x() + bar.get_width() / 2., height + 0.1,
                         f'{int(height)}', ha='center', va='bottom', fontsize=10)

            # Второй график - по уровням (круговая диаграмма)
            if level_data:
                levels = [item[0] for item in level_data]
                level_counts = [item[1] for item in level_data]

                colors2 = plt.cm.Pastel1(np.linspace(0, 1, len(levels)))

                # Круговая диаграмма вместо столбчатой
                wedges, texts, autotexts = ax2.pie(level_counts, labels=levels, autopct='%1.1f%%',
                                                   colors=colors2, startangle=90,
                                                   textprops={'fontsize': 11})
                ax2.set_title('Распределение по уровням', fontsize=14, fontweight='bold')

                # Делаем процентные метки жирными
                for autotext in autotexts:
                    autotext.set_color('black')
                    autotext.set_fontweight('bold')

            plt.suptitle('Статистика учебных достижений', fontsize=16, fontweight='bold', y=1.02)
            plt.tight_layout()

            # Встраивание графика в Tkinter
            canvas = FigureCanvasTkAgg(fig, master=chart_window)
            canvas.draw()
            canvas.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)

            # Панель управления
            control_frame = ttk.Frame(chart_window)
            control_frame.pack(pady=10)

            def save_chart():
                filename = filedialog.asksaveasfilename(
                    defaultextension=".png",
                    filetypes=[
                        ("PNG files", "*.png"),
                        ("PDF files", "*.pdf"),
                        ("SVG files", "*.svg"),
                        ("All files", "*.*")
                    ]
                )
                if filename:
                    try:
                        fig.savefig(filename, dpi=300, bbox_inches='tight')
                        messagebox.showinfo("Сохранено", f"График сохранен как: {filename}")
                    except Exception as e:
                        messagebox.showerror("Ошибка", f"Не удалось сохранить: {str(e)}")

            ttk.Button(control_frame, text="💾 Сохранить график",
                       command=save_chart, style='Accent.TButton').pack(side='left', padx=5)

            ttk.Button(control_frame, text="❌ Закрыть",
                       command=chart_window.destroy).pack(side='left', padx=5)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось построить графики: {str(e)}")

    def perform_search(self):
        """Выполнение поиска по критериям"""
        # Очистка предыдущих результатов
        for item in self.search_tree.get_children():
            self.search_tree.delete(item)

        # Получение параметров поиска
        name_filter = self.search_name.get().strip()
        type_filter = self.search_type.get()
        level_filter = self.search_level.get()
        date_from = self.date_from.get()
        date_to = self.date_to.get()

        # Построение SQL запроса
        query = "SELECT дата, название, тип, уровень, описание FROM достижения WHERE 1=1"
        params = []

        if name_filter:
            query += " AND название LIKE ?"
            params.append(f"%{name_filter}%")

        if type_filter != "Все":
            query += " AND тип = ?"
            params.append(type_filter)

        if level_filter != "Все":
            query += " AND уровень = ?"
            params.append(level_filter)

        if date_from:
            query += " AND дата >= ?"
            params.append(date_from)

        if date_to:
            query += " AND дата <= ?"
            params.append(date_to)

        query += " ORDER BY дата DESC"

        # Выполнение запроса
        try:
            self.cursor.execute(query, params)
            results = self.cursor.fetchall()

            # Отображение результатов
            for record in results:
                # Обработка None значений
                processed_record = []
                for value in record:
                    if value is None:
                        processed_record.append("")
                    else:
                        processed_record.append(str(value))
                self.search_tree.insert('', 'end', values=processed_record)

            # Показ количества найденных записей
            count_label = ttk.Label(self.tab_search,
                                    text=f"Найдено записей: {len(results)}",
                                    font=('Arial', 10, 'bold'))
            count_label.place(x=20, y=280)

            # Удаление старой метки если есть
            if hasattr(self, 'search_count_label'):
                self.search_count_label.destroy()
            self.search_count_label = count_label

        except Exception as e:
            messagebox.showerror("Ошибка поиска", f"Не удалось выполнить поиск: {str(e)}")

    def reset_search(self):
        """Сброс параметров поиска"""
        self.search_name.delete(0, tk.END)
        self.search_type.current(0)
        self.search_level.current(0)
        self.date_from.set_date(datetime.now().replace(day=1))
        self.date_to.set_date(datetime.now())

        # Очистка результатов
        for item in self.search_tree.get_children():
            self.search_tree.delete(item)

        # Удаление метки количества
        if hasattr(self, 'search_count_label'):
            self.search_count_label.destroy()

    def check_notifications(self):
        """Проверка уведомлений"""
        # Здесь можно реализовать проверку предстоящих событий
        pass

    def start_notification_scheduler(self):
        """Запуск планировщика уведомлений в отдельном потоке"""

        def scheduler_thread():
            # Проверка каждые 10 минут
            schedule.every(10).minutes.do(self.check_notifications)

            while True:
                schedule.run_pending()
                time.sleep(1)

        thread = threading.Thread(target=scheduler_thread, daemon=True)
        thread.start()

    def backup_database(self):
        """Создание резервной копии базы данных"""
        try:
            if not os.path.exists("достижения.db"):
                messagebox.showwarning("Нет данных", "База данных не найдена")
                return

            backup_name = f"достижения_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db"

            # Копируем файл
            shutil.copy2("достижения.db", backup_name)

            messagebox.showinfo("Резервная копия",
                                f"База данных успешно сохранена как:\n{backup_name}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать резервную копию: {str(e)}")

    def run(self):
        """Запуск приложения"""
        # Создание меню
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # Меню "Файл"
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Файл", menu=file_menu)
        file_menu.add_command(label="Создать резервную копию", command=self.backup_database)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self.root.quit)

        # Меню "Справка"
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Справка", menu=help_menu)
        help_menu.add_command(label="О программе",
                              command=lambda: messagebox.showinfo("О программе",
                                                                  "Журнал личных учебных достижений\nВерсия 2.0\n\nФункции:\n"
                                                                  "- Добавление и редактирование достижений\n"
                                                                  "- Поиск и фильтрация\n"
                                                                  "- Статистика и графики\n"
                                                                  "- Экспорт в Word, Excel, PDF\n"
                                                                  "- Резервное копирование"))

        # Запуск главного цикла
        self.root.mainloop()


if __name__ == "__main__":
    # Создание файла types.json если его нет
    if not os.path.exists("types.json"):
        default_types = ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция",
                         "Курс", "Публикация", "Патенты", "Хакатон", "Мастер-класс"]
        with open("types.json", "w", encoding="utf-8") as f:
            json.dump(default_types, f, ensure_ascii=False, indent=2)

    # Запуск приложения
    app = AchievementTracker()
    app.run()