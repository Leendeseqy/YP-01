import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import psycopg2
from psycopg2 import Error
import os
from datetime import datetime, timedelta
import markdown
import openpyxl
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.drawing.image import Image as XLImage
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import sys
import traceback
from PIL import Image as PILImage

# Конфигурация базы данных
DB_CONFIG = {
    'host': 'localhost',
    'database': 'postgres',
    'user': 'postgres',
    'password': '1111',
    'port': '5432'
}


class ProjectManagerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Менеджер проектов")
        self.root.geometry("1200x800")

        # Текущий выбранный проект
        self.current_project_id = None
        self.current_project_file = None

        # Словарь для хранения технологий проекта
        self.project_technologies = {}

        # Создаем структуру БД при запуске
        self.init_database()

        # Создаем папки для хранения данных
        self.create_folders()

        # Строим интерфейс
        self.setup_ui()

        # Загружаем проекты
        self.load_projects()

        # Обработка закрытия окна
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def create_folders(self):
        """Создание необходимых папок"""
        folders = ['projects', 'reports', 'reports/charts']
        for folder in folders:
            if not os.path.exists(folder):
                os.makedirs(folder)

    def init_database(self):
        """Инициализация структуры базы данных"""
        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Таблица проектов
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS projects (
                    id SERIAL PRIMARY KEY,
                    name VARCHAR(255) NOT NULL,
                    discipline VARCHAR(255),
                    status VARCHAR(100),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    file_path TEXT
                )
            """)

            # Таблица технологий
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS technologies (
                    id SERIAL PRIMARY KEY,
                    project_id INTEGER REFERENCES projects(id) ON DELETE CASCADE,
                    technology VARCHAR(255) NOT NULL,
                    added_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Таблица логов действий
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS activity_log (
                    id SERIAL PRIMARY KEY,
                    project_id INTEGER REFERENCES projects(id) ON DELETE CASCADE,
                    action_type VARCHAR(50) NOT NULL,
                    action_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    details TEXT
                )
            """)

            conn.commit()
            cursor.close()
            conn.close()

        except Error as e:
            messagebox.showerror("Ошибка БД", f"Не удалось инициализировать БД:\n{str(e)}")
            self.root.destroy()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        # Основной контейнер
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 1. Панель управления сверху
        control_frame = ttk.LabelFrame(main_frame, text="Панель управления", padding=10)
        control_frame.pack(fill=tk.X, pady=(0, 10))

        # Поля ввода
        ttk.Label(control_frame, text="Название проекта:").grid(row=0, column=0, padx=5, pady=5)
        self.project_name_entry = ttk.Entry(control_frame, width=30)
        self.project_name_entry.grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(control_frame, text="Дисциплина:").grid(row=0, column=2, padx=5, pady=5)
        self.discipline_entry = ttk.Entry(control_frame, width=20)
        self.discipline_entry.grid(row=0, column=3, padx=5, pady=5)

        ttk.Label(control_frame, text="Статус:").grid(row=0, column=4, padx=5, pady=5)
        self.status_combobox = ttk.Combobox(control_frame, width=15,
                                            values=["В процессе", "Завершен", "На паузе", "Планируется"])
        self.status_combobox.grid(row=0, column=5, padx=5, pady=5)
        self.status_combobox.set("В процессе")

        # Кнопки управления
        buttons_frame = ttk.Frame(control_frame)
        buttons_frame.grid(row=0, column=6, padx=20, pady=5)

        ttk.Button(buttons_frame, text="Создать", command=self.create_project).pack(side=tk.LEFT, padx=2)
        ttk.Button(buttons_frame, text="Сохранить", command=self.save_project).pack(side=tk.LEFT, padx=2)
        ttk.Button(buttons_frame, text="Удалить", command=self.delete_project).pack(side=tk.LEFT, padx=2)
        ttk.Button(buttons_frame, text="Открыть описание", command=self.open_description).pack(side=tk.LEFT, padx=2)
        ttk.Button(buttons_frame, text="Экспорт в Excel", command=self.export_to_excel).pack(side=tk.LEFT, padx=2)
        ttk.Button(buttons_frame, text="Экспорт в Word", command=self.export_to_word).pack(side=tk.LEFT, padx=2)

        # 2. Основная область (список проектов и редактор)
        content_frame = ttk.Frame(main_frame)
        content_frame.pack(fill=tk.BOTH, expand=True)

        # Список проектов слева
        list_frame = ttk.LabelFrame(content_frame, text="Список проектов", padding=10)
        list_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Treeview для проектов
        columns = ("Название", "Дисциплина", "Статус", "Дата создания", "Дата обновления")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=20)

        # Настройка колонок
        for col in columns:
            self.tree.heading(col, text=col, command=lambda c=col: self.sort_treeview(c))
            self.tree.column(col, width=150)

        # Полоса прокрутки
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Привязка события выбора
        self.tree.bind("<<TreeviewSelect>>", self.on_project_select)

        # 3. Область редактирования справа
        editor_frame = ttk.LabelFrame(content_frame, text="Редактирование описания", padding=10)
        editor_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # Подсказки по синтаксису
        help_frame = ttk.Frame(editor_frame)
        help_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(help_frame, text="Подсказки: ").pack(side=tk.LEFT)
        ttk.Label(help_frame, text="**жирный**", font=("TkDefaultFont", 10, "bold")).pack(side=tk.LEFT, padx=5)
        ttk.Label(help_frame, text="*курсив*", font=("TkDefaultFont", 10, "italic")).pack(side=tk.LEFT, padx=5)
        ttk.Label(help_frame, text="# Заголовок", font=("TkDefaultFont", 10, "bold")).pack(side=tk.LEFT, padx=5)
        ttk.Label(help_frame, text="- список", font=("TkDefaultFont", 10)).pack(side=tk.LEFT, padx=5)

        # Текстовый редактор
        text_frame = ttk.Frame(editor_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)

        self.text_editor = tk.Text(text_frame, wrap=tk.WORD, font=("Courier New", 11))
        text_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=self.text_editor.yview)
        self.text_editor.configure(yscrollcommand=text_scrollbar.set)

        self.text_editor.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 4. Панель технологий снизу
        tech_frame = ttk.LabelFrame(main_frame, text="Технологии проекта", padding=10)
        tech_frame.pack(fill=tk.X, pady=(10, 0))

        # Поле ввода и кнопка для технологий
        input_frame = ttk.Frame(tech_frame)
        input_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(input_frame, text="Новая технология:").pack(side=tk.LEFT, padx=(0, 5))
        self.tech_entry = ttk.Entry(input_frame, width=30)
        self.tech_entry.pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(input_frame, text="Добавить технологию", command=self.add_technology).pack(side=tk.LEFT)

        # Область отображения технологий
        self.tech_display_frame = ttk.Frame(tech_frame)
        self.tech_display_frame.pack(fill=tk.BOTH, expand=True)

        # 5. Вкладка "Аналитика и отчётность"
        self.setup_analytics_tab(main_frame)

    def setup_analytics_tab(self, parent):
        """Настройка вкладки аналитики"""
        analytics_frame = ttk.LabelFrame(parent, text="Аналитика и отчётность", padding=10)
        analytics_frame.pack(fill=tk.X, pady=(10, 0))

        ttk.Button(analytics_frame, text="Сформировать отчёт",
                   command=self.generate_report, width=30).pack(pady=10)

        # Область для вывода информации о отчете
        self.report_info_label = ttk.Label(analytics_frame, text="", foreground="green")
        self.report_info_label.pack()

    def log_activity(self, project_id, action_type, details=""):
        """Логирование действий"""
        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                INSERT INTO activity_log (project_id, action_type, details)
                VALUES (%s, %s, %s)
            """, (project_id, action_type, details))

            conn.commit()
            cursor.close()
            conn.close()

        except Error as e:
            print(f"Ошибка логирования: {e}")

    def create_project(self):
        """Создание нового проекта"""
        name = self.project_name_entry.get().strip()
        discipline = self.discipline_entry.get().strip()
        status = self.status_combobox.get()

        if not name:
            messagebox.showwarning("Предупреждение", "Введите название проекта")
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Создаем файл для проекта
            safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            file_name = f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            file_path = os.path.join('projects', file_name)

            # Создаем пустой файл
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"# {name}\n\nОписание проекта...")

            # Сохраняем в БД
            cursor.execute("""
                INSERT INTO projects (name, discipline, status, file_path)
                VALUES (%s, %s, %s, %s) RETURNING id
            """, (name, discipline, status, file_path))

            project_id = cursor.fetchone()[0]

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(project_id, "CREATE", f"Создан проект: {name}")

            # Обновляем список
            self.load_projects()

            # Очищаем поля
            self.project_name_entry.delete(0, tk.END)
            self.discipline_entry.delete(0, tk.END)

            messagebox.showinfo("Успех", f"Проект '{name}' успешно создан!")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось создать проект:\n{str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Неожиданная ошибка:\n{str(e)}")

    def load_projects(self):
        """Загрузка проектов из БД в Treeview"""
        # Очищаем текущий список
        for item in self.tree.get_children():
            self.tree.delete(item)

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT id, name, discipline, status, 
                       created_at, updated_at 
                FROM projects ORDER BY updated_at DESC
            """)

            projects = cursor.fetchall()

            for project in projects:
                # Форматируем даты
                created = project[4].strftime('%Y-%m-%d %H:%M') if project[4] else ''
                updated = project[5].strftime('%Y-%m-%d %H:%M') if project[5] else ''

                self.tree.insert("", tk.END, values=(
                    project[1],  # name
                    project[2],  # discipline
                    project[3],  # status
                    created,
                    updated
                ), tags=(f"id_{project[0]}",))

            cursor.close()
            conn.close()

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить проекты:\n{str(e)}")

    def on_project_select(self, event):
        """Обработка выбора проекта"""
        try:
            selection = self.tree.selection()
            if not selection:
                return

            item = self.tree.item(selection[0])
            project_name = item['values'][0]

            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Получаем ID проекта и путь к файлу
            cursor.execute("""
                SELECT id, file_path FROM projects WHERE name = %s
            """, (project_name,))

            result = cursor.fetchone()
            if result:
                self.current_project_id = result[0]
                self.current_project_file = result[1]

                # Загружаем описание из файла
                if os.path.exists(self.current_project_file):
                    with open(self.current_project_file, 'r', encoding='utf-8') as f:
                        content = f.read()

                    self.text_editor.delete(1.0, tk.END)
                    self.text_editor.insert(1.0, content)

                # Загружаем технологии
                self.load_technologies()

                # Обновляем поля ввода
                cursor.execute("""
                    SELECT name, discipline, status FROM projects WHERE id = %s
                """, (self.current_project_id,))

                proj_data = cursor.fetchone()
                if proj_data:
                    self.project_name_entry.delete(0, tk.END)
                    self.project_name_entry.insert(0, proj_data[0])
                    self.discipline_entry.delete(0, tk.END)
                    self.discipline_entry.insert(0, proj_data[1])
                    self.status_combobox.set(proj_data[2])

            cursor.close()
            conn.close()

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить проект:\n{str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка чтения файла:\n{str(e)}")

    def save_project(self):
        """Сохранение изменений проекта"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение", "Выберите проект для сохранения")
            return

        name = self.project_name_entry.get().strip()
        discipline = self.discipline_entry.get().strip()
        status = self.status_combobox.get()

        if not name:
            messagebox.showwarning("Предупреждение", "Введите название проекта")
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Обновляем данные в БД
            cursor.execute("""
                UPDATE projects 
                SET name = %s, discipline = %s, status = %s, updated_at = CURRENT_TIMESTAMP
                WHERE id = %s
            """, (name, discipline, status, self.current_project_id))

            # Сохраняем описание в файл
            content = self.text_editor.get(1.0, tk.END)
            if self.current_project_file and os.path.exists(self.current_project_file):
                with open(self.current_project_file, 'w', encoding='utf-8') as f:
                    f.write(content)

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(self.current_project_id, "UPDATE", f"Обновлен проект: {name}")

            # Обновляем список
            self.load_projects()

            messagebox.showinfo("Успех", "Проект успешно сохранен!")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить проект:\n{str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка записи файла:\n{str(e)}")

    def delete_project(self):
        """Удаление проекта"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение", "Выберите проект для удаления")
            return

        if not messagebox.askyesno("Подтверждение", "Вы уверены, что хотите удалить проект?"):
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Получаем информацию о проекте для логирования
            cursor.execute("SELECT name FROM projects WHERE id = %s", (self.current_project_id,))
            project_name = cursor.fetchone()[0]

            # Получаем путь к файлу
            cursor.execute("SELECT file_path FROM projects WHERE id = %s", (self.current_project_id,))
            file_path = cursor.fetchone()[0]

            # Удаляем из БД (каскадно удалятся технологии и логи)
            cursor.execute("DELETE FROM projects WHERE id = %s", (self.current_project_id,))

            # Удаляем файл
            if file_path and os.path.exists(file_path):
                os.remove(file_path)

            conn.commit()
            cursor.close()
            conn.close()

            # Очищаем интерфейс
            self.current_project_id = None
            self.current_project_file = None
            self.text_editor.delete(1.0, tk.END)
            self.project_name_entry.delete(0, tk.END)
            self.discipline_entry.delete(0, tk.END)
            self.clear_technologies_display()

            # Обновляем список
            self.load_projects()

            messagebox.showinfo("Успех", f"Проект '{project_name}' удален")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить проект:\n{str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при удалении:\n{str(e)}")

    def open_description(self):
        """Открытие описания во внешнем просмотрщике"""
        if not self.current_project_file or not os.path.exists(self.current_project_file):
            messagebox.showwarning("Предупреждение", "Файл описания не найден")
            return

        try:
            # Пытаемся открыть файл в системном просмотрщике
            if sys.platform == "win32":
                os.startfile(self.current_project_file)
            elif sys.platform == "darwin":
                os.system(f"open '{self.current_project_file}'")
            else:
                os.system(f"xdg-open '{self.current_project_file}'")
        except:
            messagebox.showinfo("Информация",
                                f"Файл расположен по пути:\n{os.path.abspath(self.current_project_file)}")

    def add_technology(self):
        """Добавление технологии к проекту"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение", "Выберите проект для добавления технологии")
            return

        tech = self.tech_entry.get().strip()
        if not tech:
            messagebox.showwarning("Предупреждение", "Введите название технологии")
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Проверяем, есть ли уже такая технология у проекта
            cursor.execute("""
                SELECT id FROM technologies 
                WHERE project_id = %s AND LOWER(technology) = LOWER(%s)
            """, (self.current_project_id, tech))

            if cursor.fetchone():
                messagebox.showwarning("Предупреждение", "Эта технология уже добавлена к проекту")
                return

            # Добавляем технологию
            cursor.execute("""
                INSERT INTO technologies (project_id, technology)
                VALUES (%s, %s)
            """, (self.current_project_id, tech))

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(self.current_project_id, "ADD_TECH", f"Добавлена технология: {tech}")

            # Обновляем отображение
            self.load_technologies()

            # Очищаем поле ввода
            self.tech_entry.delete(0, tk.END)

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить технологию:\n{str(e)}")

    def load_technologies(self):
        """Загрузка технологий проекта"""
        if not self.current_project_id:
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT technology FROM technologies 
                WHERE project_id = %s 
                ORDER BY added_at DESC
            """, (self.current_project_id,))

            technologies = [row[0] for row in cursor.fetchall()]

            cursor.close()
            conn.close()

            # Отображаем технологии
            self.display_technologies(technologies)

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить технологии:\n{str(e)}")

    def display_technologies(self, technologies):
        """Отображение технологий в интерфейсе"""
        # Очищаем текущее отображение
        self.clear_technologies_display()

        if not technologies:
            label = ttk.Label(self.tech_display_frame, text="Технологии не добавлены",
                              foreground="gray")
            label.pack(pady=10)
            return

        # Создаем фрейм с переносом по словам
        for i, tech in enumerate(technologies):
            frame = ttk.Frame(self.tech_display_frame)
            frame.pack(fill=tk.X, pady=2)

            label = ttk.Label(frame, text=f"• {tech}",
                              relief=tk.RIDGE, padding=5, width=30)
            label.pack(side=tk.LEFT, padx=(0, 5))

            # Кнопка удаления
            btn = ttk.Button(frame, text="Удалить", width=10,
                             command=lambda t=tech: self.remove_technology(t))
            btn.pack(side=tk.RIGHT)

    def clear_technologies_display(self):
        """Очистка области отображения технологий"""
        for widget in self.tech_display_frame.winfo_children():
            widget.destroy()

    def remove_technology(self, technology):
        """Удаление технологии из проекта"""
        if not self.current_project_id:
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                DELETE FROM technologies 
                WHERE project_id = %s AND technology = %s
            """, (self.current_project_id, technology))

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(self.current_project_id, "REMOVE_TECH", f"Удалена технология: {technology}")

            # Обновляем отображение
            self.load_technologies()

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить технологию:\n{str(e)}")

    def sort_treeview(self, column):
        """Сортировка Treeview по колонке"""
        # Получаем текущие данные
        items = [(self.tree.set(child, column), child) for child in self.tree.get_children('')]

        # Определяем тип сортировки
        try:
            items.sort(key=lambda x: datetime.strptime(x[0], '%Y-%m-%d %H:%M')
            if 'Дата' in column else x[0])
        except:
            items.sort()

        # Перестраиваем Treeview
        for index, (_, child) in enumerate(items):
            self.tree.move(child, '', index)

    def generate_report(self):
        """Генерация комплексного отчета"""
        try:
            # Проверяем, есть ли данные
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM projects")
            count = cursor.fetchone()[0]
            cursor.close()
            conn.close()

            if count == 0:
                messagebox.showwarning("Нет данных", "Нет проектов для генерации отчета")
                return

            # Собираем данные
            stats = self.collect_statistics()

            # Генерируем Excel отчет
            excel_path = self.generate_excel_report(stats)

            # Генерируем Word отчет
            word_path = self.generate_word_report(stats)

            # Создаем графики для Word
            self.create_charts_for_word(stats)

            # Обновляем информацию в интерфейсе
            self.report_info_label.config(
                text=f"Отчеты сгенерированы:\n{os.path.basename(excel_path)}\n{os.path.basename(word_path)}"
            )

            messagebox.showinfo("Успех", "Отчеты успешно сгенерированы!")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сгенерировать отчет:\n{str(e)}")
            print(traceback.format_exc())

    def collect_statistics(self):
        """Сбор статистических данных"""
        stats = {}

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # 1. Количество проектов по дисциплинам
            cursor.execute("""
                SELECT discipline, COUNT(*) 
                FROM projects 
                WHERE discipline IS NOT NULL AND discipline != ''
                GROUP BY discipline 
                ORDER BY COUNT(*) DESC
            """)
            stats['projects_by_discipline'] = dict(cursor.fetchall())

            # 2. Количество проектов по статусам
            cursor.execute("""
                SELECT status, COUNT(*) 
                FROM projects 
                GROUP BY status 
                ORDER BY COUNT(*) DESC
            """)
            stats['projects_by_status'] = dict(cursor.fetchall())

            # 3. Действия за последние 7 и 30 дней
            cursor.execute("""
                SELECT 
                    COUNT(CASE WHEN action_date >= CURRENT_DATE - INTERVAL '7 days' THEN 1 END) as last_7_days,
                    COUNT(CASE WHEN action_date >= CURRENT_DATE - INTERVAL '30 days' THEN 1 END) as last_30_days
                FROM activity_log
            """)
            actions = cursor.fetchone()
            stats['actions_last_7_days'] = actions[0] if actions else 0
            stats['actions_last_30_days'] = actions[1] if actions else 0

            # 4. Топ-5 самых часто используемых технологий
            cursor.execute("""
                SELECT technology, COUNT(*) as usage_count
                FROM technologies
                GROUP BY technology
                ORDER BY usage_count DESC
                LIMIT 5
            """)
            stats['top_technologies'] = dict(cursor.fetchall())

            # 5. Последние 5 проектов
            cursor.execute("""
                SELECT name, discipline, status, updated_at
                FROM projects
                ORDER BY updated_at DESC
                LIMIT 5
            """)
            stats['recent_projects'] = cursor.fetchall()

            # 6. Общее количество проектов
            cursor.execute("SELECT COUNT(*) FROM projects")
            stats['total_projects'] = cursor.fetchone()[0]

            # 7. Количество дисциплин
            stats['disciplines_count'] = len(stats['projects_by_discipline'])

            cursor.close()
            conn.close()

        except Error as e:
            raise Exception(f"Ошибка сбора статистики: {str(e)}")

        return stats

    def generate_excel_report(self, stats):
        """Генерация Excel отчета с графиками"""
        try:
            wb = Workbook()

            # Лист "Статистика"
            ws_stats = wb.active
            ws_stats.title = "Статистика"

            # Заголовок
            ws_stats['A1'] = "Отчет по проектам"
            ws_stats['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws_stats['A2'] = f"Сформирован: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

            # Основные метрики
            ws_stats['A4'] = "Основные метрики:"
            ws_stats['A4'].font = openpyxl.styles.Font(bold=True)

            data_rows = [
                ("Общее количество проектов", stats['total_projects']),
                ("Действий за 7 дней", stats['actions_last_7_days']),
                ("Действий за 30 дней", stats['actions_last_30_days']),
                ("Количество дисциплин", stats['disciplines_count']),
            ]

            for i, (label, value) in enumerate(data_rows, start=5):
                ws_stats[f'A{i}'] = label
                ws_stats[f'B{i}'] = value

            # Проекты по дисциплинам
            start_row = len(data_rows) + 7
            ws_stats[f'A{start_row}'] = "Проекты по дисциплинам:"
            ws_stats[f'A{start_row}'].font = openpyxl.styles.Font(bold=True)

            for i, (discipline, count) in enumerate(stats['projects_by_discipline'].items(), start=1):
                ws_stats[f'A{start_row + i}'] = discipline
                ws_stats[f'B{start_row + i}'] = count

            # Проекты по статусам
            start_row += len(stats['projects_by_discipline']) + 3
            ws_stats[f'A{start_row}'] = "Проекты по статусам:"
            ws_stats[f'A{start_row}'].font = openpyxl.styles.Font(bold=True)

            for i, (status, count) in enumerate(stats['projects_by_status'].items(), start=1):
                ws_stats[f'A{start_row + i}'] = status
                ws_stats[f'B{start_row + i}'] = count

            # Топ технологий
            start_row += len(stats['projects_by_status']) + 3
            ws_stats[f'A{start_row}'] = "Топ-5 технологий:"
            ws_stats[f'A{start_row}'].font = openpyxl.styles.Font(bold=True)

            for i, (tech, count) in enumerate(stats['top_technologies'].items(), start=1):
                ws_stats[f'A{start_row + i}'] = tech
                ws_stats[f'B{start_row + i}'] = count

            # Лист "Графики" - создаем встроенные графики
            ws_charts = wb.create_sheet("Графики")
            self.create_excel_charts(stats, ws_charts)

            # Сохраняем файл
            report_dir = 'reports'
            if not os.path.exists(report_dir):
                os.makedirs(report_dir)

            excel_path = os.path.join(report_dir, f'projects_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx')
            wb.save(excel_path)

            return os.path.abspath(excel_path)

        except Exception as e:
            print(f"Ошибка при генерации Excel отчета: {e}")
            # Пробуем альтернативный метод без графиков
            return self.generate_excel_simple(stats)

    def create_excel_charts(self, stats, ws_charts):
        """Создание встроенных графиков в Excel"""
        try:
            # 1. График по статусам
            ws_charts['A1'] = 'Статусы проектов'
            ws_charts['A3'] = 'Статус'
            ws_charts['B3'] = 'Количество'

            row = 4
            for status, count in stats['projects_by_status'].items():
                ws_charts[f'A{row}'] = status
                ws_charts[f'B{row}'] = count
                row += 1

            chart1 = BarChart()
            chart1.type = "col"
            chart1.style = 10
            chart1.title = "Распределение проектов по статусам"
            chart1.y_axis.title = "Количество проектов"
            chart1.x_axis.title = "Статус"
            chart1.legend = None

            data1 = Reference(ws_charts, min_col=2, min_row=3, max_row=row - 1)
            cats1 = Reference(ws_charts, min_col=1, min_row=4, max_row=row - 1)
            chart1.add_data(data1, titles_from_data=True)
            chart1.set_categories(cats1)

            ws_charts.add_chart(chart1, "D2")

            # 2. График по дисциплинам
            ws_charts[f'A{row + 2}'] = 'Проекты по дисциплинам'
            ws_charts[f'A{row + 4}'] = 'Дисциплина'
            ws_charts[f'B{row + 4}'] = 'Количество'

            row2 = row + 5
            for discipline, count in stats['projects_by_discipline'].items():
                ws_charts[f'A{row2}'] = discipline
                ws_charts[f'B{row2}'] = count
                row2 += 1

            chart2 = BarChart()
            chart2.type = "col"
            chart2.style = 11
            chart2.title = "Распределение проектов по дисциплинам"
            chart2.y_axis.title = "Количество проектов"
            chart2.x_axis.title = "Дисциплина"
            chart2.legend = None

            data2 = Reference(ws_charts, min_col=2, min_row=row + 4, max_row=row2 - 1)
            cats2 = Reference(ws_charts, min_col=1, min_row=row + 5, max_row=row2 - 1)
            chart2.add_data(data2, titles_from_data=True)
            chart2.set_categories(cats2)

            ws_charts.add_chart(chart2, "D20")

        except Exception as e:
            print(f"Ошибка при создании встроенных графиков: {e}")
            # Создаем графики как изображения
            self.create_excel_charts_as_images(stats, ws_charts)

    def create_excel_charts_as_images(self, stats, worksheet):
        """Создание графиков как изображений для Excel"""
        try:
            charts_dir = 'reports/charts'
            if not os.path.exists(charts_dir):
                os.makedirs(charts_dir)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # 1. График проектов по статусам
            if stats['projects_by_status']:
                fig, ax = plt.subplots(figsize=(6, 4))
                labels = list(stats['projects_by_status'].keys())
                values = list(stats['projects_by_status'].values())

                bars = ax.bar(labels, values, color='lightblue')
                ax.set_title('Проекты по статусам', fontsize=10)
                ax.set_ylabel('Количество', fontsize=8)

                # Добавляем значения на столбцы
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.05,
                            f'{int(height)}', ha='center', va='bottom', fontsize=8)

                plt.xticks(rotation=45, ha='right', fontsize=8)
                plt.tight_layout()

                chart_path = os.path.join(charts_dir, f'excel_status_{timestamp}.png')
                plt.savefig(chart_path, dpi=100, bbox_inches='tight')
                plt.close()

                # Уменьшаем размер изображения
                try:
                    img = PILImage.open(chart_path)
                    img.thumbnail((400, 300))
                    img.save(chart_path, 'PNG')

                    # Вставляем в Excel
                    excel_img = XLImage(chart_path)
                    worksheet.add_image(excel_img, 'A1')
                except:
                    print("Не удалось вставить изображение в Excel")

            # 2. График проектов по дисциплинам
            if stats['projects_by_discipline']:
                fig, ax = plt.subplots(figsize=(6, 4))
                labels = list(stats['projects_by_discipline'].keys())
                values = list(stats['projects_by_discipline'].values())

                bars = ax.bar(labels, values, color='lightgreen')
                ax.set_title('Проекты по дисциплинам', fontsize=10)
                ax.set_ylabel('Количество', fontsize=8)

                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.05,
                            f'{int(height)}', ha='center', va='bottom', fontsize=8)

                plt.xticks(rotation=45, ha='right', fontsize=8)
                plt.tight_layout()

                chart_path = os.path.join(charts_dir, f'excel_discipline_{timestamp}.png')
                plt.savefig(chart_path, dpi=100, bbox_inches='tight')
                plt.close()

                # Уменьшаем размер изображения
                try:
                    img = PILImage.open(chart_path)
                    img.thumbnail((400, 300))
                    img.save(chart_path, 'PNG')

                    excel_img = XLImage(chart_path)
                    worksheet.add_image(excel_img, 'A20')
                except:
                    print("Не удалось вставить изображение в Excel")

        except Exception as e:
            print(f"Ошибка при создании изображений для Excel: {e}")

    def create_charts_for_word(self, stats):
        """Создание графиков для Word отчета"""
        try:
            charts_dir = 'reports/charts'
            if not os.path.exists(charts_dir):
                os.makedirs(charts_dir)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # График проектов по статусам для Word
            if stats['projects_by_status']:
                fig, ax = plt.subplots(figsize=(8, 5))
                labels = list(stats['projects_by_status'].keys())
                values = list(stats['projects_by_status'].values())

                colors = ['#4CAF50', '#2196F3', '#FF9800', '#F44336']
                bars = ax.bar(labels, values, color=colors[:len(labels)])
                ax.set_title('Распределение проектов по статусам', fontsize=12, fontweight='bold')
                ax.set_ylabel('Количество проектов', fontsize=10)
                ax.set_xlabel('Статус', fontsize=10)

                # Добавляем значения на столбцы
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.1,
                            f'{int(height)}', ha='center', va='bottom', fontsize=10, fontweight='bold')

                plt.xticks(fontsize=9)
                plt.tight_layout()

                chart_path = os.path.join(charts_dir, f'word_status_{timestamp}.png')
                plt.savefig(chart_path, dpi=150, bbox_inches='tight')
                plt.close()

        except Exception as e:
            print(f"Ошибка при создании графиков для Word: {e}")

    def generate_excel_simple(self, stats):
        """Альтернативная генерация Excel без графиков"""
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Статистика"

            ws['A1'] = "Отчет по проектам"
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws['A2'] = f"Сформирован: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            ws['A2'].font = openpyxl.styles.Font(italic=True)

            row = 4
            ws[f'A{row}'] = "Общее количество проектов:"
            ws[f'B{row}'] = stats['total_projects']
            row += 1

            ws[f'A{row}'] = "Действий за 7 дней:"
            ws[f'B{row}'] = stats['actions_last_7_days']
            row += 1

            ws[f'A{row}'] = "Действий за 30 дней:"
            ws[f'B{row}'] = stats['actions_last_30_days']
            row += 2

            # Проекты по статусам
            ws[f'A{row}'] = "Проекты по статусам:"
            ws[f'A{row}'].font = openpyxl.styles.Font(bold=True)
            row += 1

            for status, count in stats['projects_by_status'].items():
                ws[f'A{row}'] = f"  {status}:"
                ws[f'B{row}'] = count
                row += 1

            row += 1

            # Проекты по дисциплинам
            ws[f'A{row}'] = "Проекты по дисциплинам:"
            ws[f'A{row}'].font = openpyxl.styles.Font(bold=True)
            row += 1

            for discipline, count in stats['projects_by_discipline'].items():
                ws[f'A{row}'] = f"  {discipline}:"
                ws[f'B{row}'] = count
                row += 1

            # Сохраняем файл
            report_dir = 'reports'
            if not os.path.exists(report_dir):
                os.makedirs(report_dir)

            excel_path = os.path.join(report_dir, f'projects_simple_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx')
            wb.save(excel_path)

            return os.path.abspath(excel_path)

        except Exception as e:
            raise Exception(f"Не удалось создать Excel отчет: {e}")

    def generate_word_report(self, stats):
        """Генерация Word отчета"""
        try:
            # Создаем документ
            doc = Document()

            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.5

            # Титульный лист
            title = doc.add_heading('Отчет по проектам', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            date_para = doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()

            # Сводная таблица ключевых показателей
            doc.add_heading('Ключевые показатели', level=1)

            table = doc.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'

            # Заполняем таблицу
            data = [
                ("Общее количество проектов", str(stats['total_projects'])),
                ("Действий за 7 дней", str(stats['actions_last_7_days'])),
                ("Действий за 30 дней", str(stats['actions_last_30_days'])),
                ("Количество дисциплин", str(stats['disciplines_count'])),
                ("Количество статусов", str(len(stats['projects_by_status'])))
            ]

            for i, (label, value) in enumerate(data):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = value

            # Вставляем график
            doc.add_heading('Визуализация данных', level=1)

            # Ищем последний сохраненный график для Word
            charts_dir = 'reports/charts'
            if os.path.exists(charts_dir):
                png_files = [f for f in os.listdir(charts_dir) if f.startswith('word_status_') and f.endswith('.png')]
                if png_files:
                    latest_chart = os.path.join(charts_dir, sorted(png_files)[-1])

                    doc.add_paragraph('Распределение проектов по статусам:')
                    try:
                        doc.add_picture(latest_chart, width=Inches(6))
                    except:
                        doc.add_paragraph(f"[График доступен по пути: {latest_chart}]")

            # Последние проекты
            doc.add_heading('Последние 5 проектов', level=1)

            if stats['recent_projects']:
                projects_table = doc.add_table(rows=len(stats['recent_projects']) + 1, cols=4)
                projects_table.style = 'Light Grid Accent 1'

                # Заголовки
                headers = ["Название", "Дисциплина", "Статус", "Обновлен"]
                for i, header in enumerate(headers):
                    cell = projects_table.cell(0, i)
                    cell.text = header
                    # Делаем заголовки жирными
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Данные
                for i, project in enumerate(stats['recent_projects'], start=1):
                    for j, value in enumerate(project[:3]):
                        projects_table.cell(i, j).text = str(value) if value else ""

                    # Форматируем дату
                    if project[3]:
                        date_str = project[3].strftime('%d.%m.%Y %H:%M')
                        projects_table.cell(i, 3).text = date_str

            # Топ технологий
            if stats['top_technologies']:
                doc.add_heading('Топ используемых технологий', level=1)

                tech_table = doc.add_table(rows=len(stats['top_technologies']) + 1, cols=2)
                tech_table.style = 'Light Grid Accent 1'

                tech_table.cell(0, 0).text = "Технология"
                tech_table.cell(0, 1).text = "Использований"

                for i, (tech, count) in enumerate(stats['top_technologies'].items(), start=1):
                    tech_table.cell(i, 0).text = tech
                    tech_table.cell(i, 1).text = str(count)

            # Сохраняем документ
            report_dir = 'reports'
            if not os.path.exists(report_dir):
                os.makedirs(report_dir)

            word_path = os.path.join(report_dir, f'projects_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
            doc.save(word_path)

            return os.path.abspath(word_path)

        except Exception as e:
            raise Exception(f"Не удалось создать Word отчет: {e}")

    def export_to_excel(self):
        """Экспорт текущего списка проектов в Excel"""
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Проекты"

            # Заголовки
            headers = ["Название", "Дисциплина", "Статус", "Дата создания", "Дата обновления"]
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
                # Настраиваем ширину колонок
                ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 20

            # Данные
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT name, discipline, status, created_at, updated_at 
                FROM projects 
                ORDER BY name
            """)

            for row_idx, project in enumerate(cursor.fetchall(), start=2):
                for col_idx, value in enumerate(project, start=1):
                    if isinstance(value, datetime):
                        ws.cell(row=row_idx, column=col_idx,
                                value=value.strftime('%Y-%m-%d %H:%M'))
                    else:
                        ws.cell(row=row_idx, column=col_idx, value=value)

            cursor.close()
            conn.close()

            # Сохраняем файл
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                initialfile=f"projects_export_{datetime.now().strftime('%Y%m%d')}.xlsx"
            )

            if file_path:
                wb.save(file_path)
                messagebox.showinfo("Успех", f"Данные экспортированы в:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать данные:\n{str(e)}")

    def export_to_word(self):
        """Экспорт текущего проекта в Word"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение", "Выберите проект для экспорта")
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT name, discipline, status, created_at, updated_at 
                FROM projects 
                WHERE id = %s
            """, (self.current_project_id,))

            project = cursor.fetchone()

            # Получаем технологии
            cursor.execute("""
                SELECT technology FROM technologies 
                WHERE project_id = %s 
                ORDER BY added_at
            """, (self.current_project_id,))

            technologies = [row[0] for row in cursor.fetchall()]

            cursor.close()
            conn.close()

            # Читаем описание из файла
            content = ""
            if self.current_project_file and os.path.exists(self.current_project_file):
                with open(self.current_project_file, 'r', encoding='utf-8') as f:
                    content = f.read()

            # Создаем Word документ
            doc = Document()

            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # Заголовок
            title = doc.add_heading(project[0], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Информация о проекте
            doc.add_paragraph(f"Дисциплина: {project[1] or 'Не указана'}")
            doc.add_paragraph(f"Статус: {project[2] or 'Не указан'}")

            # Даты
            if project[3]:
                created_date = project[3].strftime('%d.%m.%Y %H:%M')
                doc.add_paragraph(f"Создан: {created_date}")

            if project[4]:
                updated_date = project[4].strftime('%d.%m.%Y %H:%M')
                doc.add_paragraph(f"Обновлен: {updated_date}")

            doc.add_paragraph()

            # Технологии
            if technologies:
                doc.add_heading('Используемые технологии', level=2)
                for tech in technologies:
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = False
                    p.add_run(tech)

            doc.add_heading('Описание проекта', level=2)

            # Конвертируем Markdown в простой текст для Word
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = False
                    p.add_run(line[2:])
                elif line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()

            # Сохраняем файл
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
                initialfile=f"{project[0].replace(' ', '_')}_report.docx"
            )

            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Успех", f"Проект экспортирован в:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать проект:\n{str(e)}")

    def on_closing(self):
        """Обработка закрытия приложения"""
        if messagebox.askokcancel("Выход", "Вы уверены, что хотите выйти?"):
            self.root.destroy()


# Запуск приложения
if __name__ == "__main__":
    # Проверка зависимостей
    required_libraries = {
        'psycopg2': 'psycopg2-binary',
        'openpyxl': 'openpyxl',
        'docx': 'python-docx',
        'matplotlib': 'matplotlib',
        'PIL': 'pillow'
    }

    missing_libs = []
    for lib, pip_name in required_libraries.items():
        try:
            __import__(lib)
        except ImportError:
            missing_libs.append(pip_name)

    if missing_libs:
        print("Ошибка: Не установлены необходимые библиотеки:")
        for lib in missing_libs:
            print(f"  - {lib}")
        print("\nУстановите зависимости:")
        print(f"pip install {' '.join(missing_libs)}")
        input("Нажмите Enter для выхода...")
        sys.exit(1)

    root = tk.Tk()
    app = ProjectManagerApp(root)
    root.mainloop()